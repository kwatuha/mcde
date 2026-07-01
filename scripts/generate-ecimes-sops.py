#!/usr/bin/env python3
"""
Generate E-CIMES Standard Operating Procedures (SOP) manual (Word).

SOPs describe official county processes — roles, approvals, records, and escalation.
The User Manual (generate-ecimes-user-manual.py) describes system features and screens.
"""

from __future__ import annotations

import shutil
import subprocess
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
LOGO_PATH = ROOT / "frontend" / "src" / "assets" / "gpris.png"
OUT_DOCX = ROOT / "docs" / "E-CIMES-Standard-Operating-Procedures.docx"
OUT_PDF = ROOT / "docs" / "E-CIMES-Standard-Operating-Procedures.pdf"

NAVY = RGBColor(0x0D, 0x47, 0xA1)
BLUE = RGBColor(0x15, 0x65, 0xC0)
GREEN = RGBColor(0x2E, 0x7D, 0x32)
SLATE = RGBColor(0x33, 0x33, 0x33)
HEADER_FILL = "0D47A1"
ALT_FILL = "E3F2FD"
SOP_VERSION = "1.0"
DOCUMENT_OWNER = "Department of Finance and Economic Planning — ICT / M&E"
REVIEW_CYCLE = "Annual, or after major system or policy change"

SOPS = [
    {
        "id": "SOP-001",
        "title": "User account request, approval, and deactivation",
        "owner": "County ICT Administrator",
        "frequency": "As required",
        "purpose": (
            "Ensure every E-CIMES user account is created with the correct role, organisation scope, "
            "and approval trail; and deactivated promptly when staff leave or change duties."
        ),
        "scope": (
            "All county staff, contractors (where applicable), and partner users who require access to "
            "E-CIMES. Excludes public citizens who use the public dashboard without login."
        ),
        "roles": [
            ("Requesting officer / HOD", "Submits account request with name, designation, department, phone, and required access."),
            ("County ICT Administrator", "Creates account, assigns role and scope, resets passwords, deactivates accounts."),
            ("Super Administrator", "Approves privileged roles (admin, finance approval, metadata changes)."),
            ("Audit / compliance", "May review user lists and audit trail on request."),
        ],
        "preconditions": [
            "Staff member is formally assigned to a county department or authorised unit.",
            "HOD or equivalent has approved the access request in writing (email or internal memo).",
            "ICT has confirmed the user does not already have an active duplicate account.",
        ],
        "steps": [
            ("Requesting officer", "Complete the county user-access request form (name, ID, designation, department, email, phone, modules needed, justification)."),
            ("HOD", "Approve the request and forward to ICT with copy to M&E focal person where M&E access is requested."),
            ("ICT Administrator", "Open Admin → User Management. Search for existing accounts before creating a new one."),
            ("ICT Administrator", "Create the account with least-privilege role and correct organisation scope (department, ward, or county-wide as approved)."),
            ("ICT Administrator", "Enable OTP if county policy requires it. Issue initial credentials through approved secure channel — never by public chat."),
            ("New user", "Log in, complete forced password change and OTP setup. Confirm visible menus match the approved role."),
            ("ICT Administrator", "On transfer or exit: deactivate the account the same day notice is received; do not delete historical audit records."),
        ],
        "records": [
            "Signed or emailed access request and HOD approval.",
            "User Management change log and Audit Trail entries for create, role change, and deactivation.",
            "Password reset tickets with requester identity verified.",
        ],
        "controls": [
            "No shared administrator accounts.",
            "Privileged roles require Super Administrator approval.",
            "Quarterly user-access review by ICT and department HODs.",
        ],
        "escalation": "Unresolved access disputes → HOD → Chief Officer (Finance & EP) → County ICT steering committee.",
        "manual_ref": "User Manual §14 User administration; §1 Accessing the system.",
    },
    {
        "id": "SOP-002",
        "title": "New project registration and ADP linkage",
        "owner": "Department focal person / Project officer",
        "frequency": "Per new project; aligned to ADP cycle",
        "purpose": (
            "Register county projects once in E-CIMES with correct planning links so dashboards, "
            "monitoring, finance, and public reporting stay consistent."
        ),
        "scope": "All new development projects, maintenance projects, and programmes tracked in the county ADP/CIDP.",
        "roles": [
            ("Department focal person", "Initiates project record with accurate metadata and ADP linkage."),
            ("Planning unit", "Validates ADP period, sector, programme, and budget line alignment."),
            ("M&E focal person", "Confirms indicators and monitoring setup after registration."),
            ("ICT Administrator", "Supports import issues and metadata corrections only — not business validation."),
        ],
        "preconditions": [
            "Project is approved in ADP or authorised supplementary allocation.",
            "CIDP/ADP catalogues for the financial year are current in E-CIMES.",
            "Project does not already exist — search Projects Registry first.",
        ],
        "steps": [
            ("Department focal person", "Open Projects → Projects Registry. Search by name, code, ward, and department to avoid duplicates."),
            ("Department focal person", "Create project with official name, department, sector, ward/subcounty, financial year, budget, funding source, and status."),
            ("Department focal person", "Link the project to ADP implementation line, CIDP programme/sub-programme, and activities where configured."),
            ("Planning unit", "Review Planning → ADP Implementation to confirm wishlist/budgeted/registered alignment."),
            ("Department focal person", "Upload inception documents (concept, approvals, site handover) on the project Documents tab."),
            ("M&E focal person", "Verify indicators, risks, and reporting frequency are linked for monitoring and APR reporting."),
            ("HOD", "Acknowledge registration complete when project appears correctly on department dashboards."),
        ],
        "records": [
            "Project record with unique identifier in Projects Registry.",
            "Linked ADP/CIDP references and uploaded approval documents.",
            "Audit Trail entry for project creation and major edits.",
        ],
        "controls": [
            "Mandatory duplicate search before create.",
            "Planning unit spot-check monthly on ADP Implementation dashboard.",
            "No project published to public dashboard until SOP-010 is followed.",
        ],
        "escalation": "Duplicate or disputed ownership → Planning + HODs → Chief Officer Planning.",
        "manual_ref": "User Manual §5 Projects Registry; §7 Planning (CIDP, ADP, traceability).",
    },
    {
        "id": "SOP-003",
        "title": "Project progress update and evidence documentation",
        "owner": "Department focal person / Project officer",
        "frequency": "Monthly minimum; immediately after major milestones",
        "purpose": (
            "Maintain accurate implementation status and documentary evidence so leadership dashboards, "
            "certificates, and county reports reflect verified progress."
        ),
        "scope": "All active projects in implementation or completion phases.",
        "roles": [
            ("Project officer", "Updates milestones, status, and progress percentages."),
            ("Site / field officer", "Captures photos and site observations."),
            ("M&E officer", "Reviews evidence quality and monitoring alignment."),
            ("HOD", "Reviews exceptions, stalled projects, and reported completion."),
        ],
        "preconditions": [
            "Project exists in Projects Registry with correct scope.",
            "Progress update is based on verified site or documentary evidence — not estimates alone.",
        ],
        "steps": [
            ("Project officer", "Open project from Projects Registry. Review current milestones and last update date."),
            ("Project officer", "Update milestone dates, completion percentages, and implementation status using verified evidence."),
            ("Project officer", "Upload supporting documents and geotagged photos with clear filenames (project code, date, subject)."),
            ("Project officer", "Add comments explaining delays, risks, or scope changes."),
            ("M&E officer", "Cross-check against Monitoring → Project Monitoring entries where visits were conducted."),
            ("HOD", "Review department Status Report monthly; flag projects with no update in 30+ days."),
            ("Project officer", "Mark practical completion only when handover/defects documentation supports it."),
        ],
        "records": [
            "Milestone and status history on project record.",
            "Documents and photos on project tabs.",
            "Monitoring visit records linked to the same project.",
        ],
        "controls": [
            "Progress percentage must match attached evidence.",
            "M&E random verification sample each quarter.",
            "Audit Trail captures status changes and document uploads.",
        ],
        "escalation": "Material misreporting or missing evidence → HOD → Chief Officer → County M&E committee.",
        "manual_ref": "User Manual §6 Project details, evaluation, and evidence.",
    },
    {
        "id": "SOP-004",
        "title": "ADP/CIDP planning maintenance and budget traceability",
        "owner": "County Planning unit",
        "frequency": "Per planning cycle; ad hoc for approved revisions",
        "purpose": (
            "Keep planning structures, budget allocations, and traceability links current so project "
            "registration and county reports align with approved plans."
        ),
        "scope": "CIDP periods, ADP periods, sectors, programmes, sub-programmes, budget lines, and indicators catalogues.",
        "roles": [
            ("Planning officer", "Maintains ADP/CIDP structures and budget allocations."),
            ("Finance officer", "Validates budget figures against approved estimates."),
            ("ICT Administrator", "Supports metadata setup and coordinates catalogue changes."),
            ("M&E focal person", "Maintains indicators, risks, and reporting frequency catalogues."),
        ],
        "preconditions": [
            "County assembly / executive approval for plan or revision is available.",
            "Change window communicated to departments before catalogue edits.",
        ],
        "steps": [
            ("Planning officer", "Open Planning → maintain CIDP period, ADP period, sectors, programmes, and sub-programmes."),
            ("Planning officer", "Enter or update budget allocations per approved ADP document."),
            ("Finance officer", "Reconcile totals with finance estimates; record discrepancies for correction."),
            ("Planning officer", "Use ADP Implementation view to track wishlist vs budgeted vs registered projects."),
            ("M&E focal person", "Update indicators, activities, risks, and reporting frequency catalogues."),
            ("Planning officer", "Use Budget Traceability to confirm budget lines flow to registered projects."),
            ("ICT Administrator", "Announce catalogue freeze dates before major reporting deadlines (APR, quarterly reviews)."),
        ],
        "records": [
            "Approved ADP/CIDP documents stored per county records policy.",
            "Planning module entries with effective financial year.",
            "Change log email to departments when catalogues are updated.",
        ],
        "controls": [
            "No catalogue change during active reporting window without Planning + M&E sign-off.",
            "Search before creating duplicate sectors/programmes.",
            "Annual reconciliation between E-CIMES totals and approved budget book.",
        ],
        "escalation": "Cross-department programme disputes → Chief Officer Planning → CEC Finance & EP.",
        "manual_ref": "User Manual §7 Planning (CIDP, ADP, RRI, traceability).",
    },
    {
        "id": "SOP-005",
        "title": "Procurement case tracking in E-CIMES",
        "owner": "Procurement unit / Department focal person",
        "frequency": "Per procurement case; updated at each stage transition",
        "purpose": (
            "Mirror approved procurement milestones in E-CIMES so project officers, finance, and "
            "leadership can track delivery against procurement timelines."
        ),
        "scope": "All county project procurements tracked through E-CIMES Procurement module.",
        "roles": [
            ("Procurement officer", "Updates procurement stages and attaches key documents."),
            ("Department focal person", "Initiates budget procurement intake and links to project."),
            ("Contractor (portal user)", "Views assigned projects on Contractor Dashboard where enabled."),
            ("HOD", "Reviews stalled procurements on operations dashboards."),
        ],
        "preconditions": [
            "Project is registered in Projects Registry.",
            "Procurement plan / requisition is approved per PPADA and county regulations.",
        ],
        "steps": [
            ("Department focal person", "Open Procurement → Budget Procurement Intake. Link item to project and approved budget line."),
            ("Procurement officer", "Open Procurement → Project Procurement. Create or open case for the project."),
            ("Procurement officer", "Update stage (requisition, advertising, evaluation, award, contract signing, etc.) to match physical file."),
            ("Procurement officer", "Attach evaluation reports, award letters, and contract summaries as procurement documents."),
            ("Procurement officer", "Record awarded contractor in Contractor Registry with correct contractor type."),
            ("Department focal person", "Confirm project procurement tab reflects award before implementation tracking intensifies."),
            ("Procurement officer", "Move completed cases to Procured Projects history when contract is active."),
        ],
        "records": [
            "Procurement stage history in E-CIMES.",
            "Scanned approvals and award documents.",
            "Contractor Registry entry linked to project.",
        ],
        "controls": [
            "Stage dates must not precede documented approval dates.",
            "Quarterly procurement audit sample by internal audit / procurement unit.",
        ],
        "escalation": "Procurement delay beyond threshold → HOD → CEC member → Procurement committee.",
        "manual_ref": "User Manual §9 Procurement and contractors.",
    },
    {
        "id": "SOP-006",
        "title": "Monitoring visit and field data collection",
        "owner": "County M&E unit",
        "frequency": "Per visit schedule; minimum monthly for priority projects",
        "purpose": (
            "Capture structured, verifiable monitoring evidence through web and mobile tools for "
            "accountability and corrective action."
        ),
        "scope": "Project monitoring visits, checklist templates, and Mobile Field Collector (Android) submissions.",
        "roles": [
            ("M&E officer", "Plans visits, reviews submissions, escalates findings."),
            ("Field enumerator", "Completes checklists on mobile app or web."),
            ("Project officer", "Responds to monitoring findings within agreed timelines."),
            ("ICT Administrator", "Publishes mobile app releases and supports sync issues."),
        ],
        "preconditions": [
            "Checklist templates are approved and synced to field devices.",
            "Field staff have active E-CIMES accounts with monitoring privileges.",
            "Mobile app is latest version from Dashboard → Mobile app (Android).",
        ],
        "steps": [
            ("M&E officer", "Open Monitoring → Checklists & visits. Confirm templates are current for the visit type."),
            ("Field enumerator", "Install/update Mobile Field Collector; sign in; pull down on Checklists tab to sync."),
            ("Field enumerator", "Select project, complete checklist (works offline); submit when connectivity allows."),
            ("M&E officer", "Open Monitoring → Project Monitoring. Record structured visit summary and link to project."),
            ("M&E officer", "Attach photos and flag delays, quality issues, or community complaints."),
            ("Project officer", "Acknowledge findings and update project status/evidence per SOP-003."),
            ("M&E officer", "Escalate unresolved critical issues to HOD and County M&E committee."),
        ],
        "records": [
            "Checklist submissions (web and mobile).",
            "Project Monitoring visit entries with timestamps.",
            "Exported visit PDFs where generated.",
        ],
        "controls": [
            "Visits must reference a valid project code.",
            "Supervisor reviews 10% of field submissions monthly.",
            "Template changes communicated before field deployment.",
        ],
        "escalation": "Safety or fraud concerns → immediate HOD + County Secretary channel per county protocol.",
        "manual_ref": "User Manual §10 Monitoring; Mobile Field Collector workflow.",
    },
    {
        "id": "SOP-007",
        "title": "PMC ward reporting",
        "owner": "Ward PMC coordinator / M&E unit",
        "frequency": "Monthly per county PMC calendar",
        "purpose": (
            "Compile ward-level Project Monitoring Committee reports for grassroots accountability "
            "and aggregation into county monitoring reports."
        ),
        "scope": "All wards with active PMC reporting requirements in E-CIMES.",
        "roles": [
            ("Ward PMC coordinator", "Completes ward PMC entries for the reporting period."),
            ("Sub-county M&E focal person", "Reviews submissions for completeness."),
            ("County M&E officer", "Consolidates ward reports for leadership review."),
        ],
        "preconditions": [
            "Reporting period and ward list are published by County M&E.",
            "Ward coordinators have accounts with access to Monitoring → PMC Ward Reports.",
        ],
        "steps": [
            ("County M&E officer", "Issue monthly PMC reporting window and ward assignment list."),
            ("Ward PMC coordinator", "Open Monitoring → PMC Ward Reports. Filter by ward and financial year."),
            ("Ward PMC coordinator", "Complete required fields for each active project in the ward."),
            ("Ward PMC coordinator", "Submit before deadline; export PDF copy for ward file if required."),
            ("Sub-county M&E focal person", "Review submissions; return incomplete entries within 2 working days."),
            ("County M&E officer", "Lock period after validation; include in county monitoring brief."),
        ],
        "records": [
            "PMC Ward Report entries per period.",
            "Exported ward PDFs filed per county records policy.",
            "County consolidated monitoring brief.",
        ],
        "controls": [
            "Late submissions tracked; repeat offenders reported to Sub-county Administrator.",
            "Cross-check ward totals against Projects Registry ward filters.",
        ],
        "escalation": "Missing ward report at deadline → Sub-county Administrator → County M&E steering.",
        "manual_ref": "User Manual §10 PMC ward reports.",
    },
    {
        "id": "SOP-008",
        "title": "Payment certificate issuance and verification",
        "owner": "Finance unit / Accounts",
        "frequency": "Per payment event (interim/final)",
        "purpose": (
            "Issue authentic, verifiable payment certificates linked to project records; enable "
            "contractors, auditors, and citizens to confirm certificate validity."
        ),
        "scope": "Interim and final payment certificates generated in E-CIMES for county projects.",
        "roles": [
            ("Project officer", "Prepares BQ/progress basis for certificate."),
            ("Quantity surveyor / engineer", "Certifies works completed where role exists."),
            ("Finance officer", "Creates certificate, runs approval workflow, generates PDF."),
            ("Approver", "Approves per delegated financial authority."),
            ("Any verifier", "Uses Finance → Verify Certificate or QR scan on PDF."),
        ],
        "preconditions": [
            "Project exists with updated progress and supporting evidence.",
            "Bill of quantities requirements are complete where configured.",
            "Approver has active finance approval privileges.",
        ],
        "steps": [
            ("Project officer", "Open project → Certificates tab. Confirm progress and BQ align with site evidence."),
            ("Finance officer", "Create interim or final certificate with correct amounts, retention, and references."),
            ("Approver", "Review and approve through workflow inbox on home page."),
            ("Finance officer", "Generate PDF. Confirm QR code appears (Scan to verify this certificate)."),
            ("Finance officer", "Verify issued PDF using Finance → Verify Certificate before external release."),
            ("Finance officer", "Issue PDF to contractor through official channel; file copy per finance records policy."),
            ("Verifier (staff/public)", "Scan QR or enter certificate number at /verify-certificate to confirm validity."),
        ],
        "records": [
            "Certificate record in Finance → Payment Certificates.",
            "Signed PDF with QR code.",
            "Verification log for disputed certificates.",
        ],
        "controls": [
            "Mandatory verification before PDF leaves finance office.",
            "Voided certificates must remain in system with void status — do not delete.",
            "Finance dashboard reviewed monthly for orphan certificates.",
        ],
        "escalation": "Disputed certificate → HOD Finance → County Auditor → if fraud suspected, EACC protocol.",
        "manual_ref": "User Manual §11 Finance and certificates; Verify certificate workflow.",
    },
    {
        "id": "SOP-009",
        "title": "Official report generation and distribution",
        "owner": "M&E unit / Report custodian",
        "frequency": "Per reporting calendar (monthly, quarterly, annual)",
        "purpose": (
            "Produce consistent, filter-documented county reports from E-CIMES for executive, "
            "assembly, and donor reporting — distinct from informal AI drafts."
        ),
        "scope": "Built-in Reports hub outputs, Report Library archives, and Scheduled Reports.",
        "roles": [
            ("Report custodian (M&E/Planning/Finance)", "Generates report with agreed filters."),
            ("HOD", "Reviews departmental figures before county consolidation."),
            ("CEC / Chief Officer", "Approves county-wide external release."),
            ("ICT Administrator", "Maintains scheduled report jobs and distribution lists."),
        ],
        "preconditions": [
            "Source data updated per SOP-003 and SOP-006 deadlines.",
            "Report template and filter conventions agreed for the period.",
        ],
        "steps": [
            ("Report custodian", "Open Reports → Reports hub. Select official template (Status, APR, Absorption, Pending Bills, etc.)."),
            ("Report custodian", "Apply standard filters (financial year, department, status). Record filter criteria in cover memo."),
            ("Report custodian", "Review totals against Projects Registry and Finance Dashboard; investigate variances."),
            ("HOD", "Sign off departmental section or email approval."),
            ("Report custodian", "Export final version; save to Report Library with period label."),
            ("ICT Administrator", "Configure Scheduled Reports for recurring internal distribution where approved."),
            ("Report custodian", "Distribute approved PDF/Excel per county communications protocol."),
        ],
        "records": [
            "Exported report file with date and filter memo.",
            "Report Library entry.",
            "Email approval from HOD/Chief Officer for external release.",
        ],
        "controls": [
            "AI Professional Reports are advisory only — not official county releases without human review.",
            "Filter criteria must accompany every shared export.",
            "Annual comparison of APR figures to planning and finance books.",
        ],
        "escalation": "Material data variance → data owner department → County M&E + Finance reconciliation meeting.",
        "manual_ref": "User Manual §12 Reports; §9 AI Assistant (advisory use only).",
    },
    {
        "id": "SOP-010",
        "title": "Public project publishing and citizen feedback",
        "owner": "Public engagement / Communications",
        "frequency": "Per project; ongoing moderation",
        "purpose": (
            "Publish accurate, non-sensitive project information to the public dashboard and manage "
            "citizen feedback responsibly."
        ),
        "scope": "Public dashboard content, Public Approval workflow, announcements, and feedback moderation.",
        "roles": [
            ("Department focal person", "Proposes projects for public visibility."),
            ("Public engagement officer", "Reviews content and moderates feedback."),
            ("HOD", "Approves sensitive projects before publishing."),
            ("ICT Administrator", "Technical support for public URL and approval workflow."),
        ],
        "preconditions": [
            "Project record is complete and verified per SOP-002 and SOP-003.",
            "No confidential procurement, personnel, or security-sensitive data in public fields.",
        ],
        "steps": [
            ("Department focal person", "Confirm project details, location, and progress are accurate."),
            ("Department focal person", "Submit project for public visibility via Public → Public Approval."),
            ("Public engagement officer", "Review for accuracy, sensitive content, and image appropriateness."),
            ("HOD", "Approve for wards/projects with political or security sensitivity."),
            ("Public engagement officer", "Approve in Public Approval; verify project appears on public dashboard."),
            ("Public engagement officer", "Monitor citizen feedback; respond or assign to department within SLA."),
            ("Public engagement officer", "Publish announcements only when approved by Communications lead."),
        ],
        "records": [
            "Public Approval audit entries.",
            "Feedback response log.",
            "Screenshot/archive of published page per communications policy.",
        ],
        "controls": [
            "No internal notes or draft figures on public-facing fields.",
            "Quarterly review of published projects for stale progress.",
        ],
        "escalation": "Sensitive or viral complaint → HOD → County Communications → County Secretary.",
        "manual_ref": "User Manual §13 Public dashboard and citizen engagement.",
    },
    {
        "id": "SOP-011",
        "title": "Bulk data import (projects, budgets, beneficiaries)",
        "owner": "ICT Administrator / Authorised data officer",
        "frequency": "As approved (migration, annual seeding)",
        "purpose": (
            "Import large datasets safely with validation preview, audit logs, and rollback plan — "
            "without corrupting production metadata."
        ),
        "scope": "Data → Import Data for projects, budgets, beneficiaries; Data Import Logs review.",
        "roles": [
            ("Authorised data officer", "Prepares and uploads spreadsheet."),
            ("Department HOD", "Approves import scope and source file."),
            ("ICT Administrator", "Runs preview, confirms import, retains logs."),
            ("M&E focal person", "Validates sample records post-import."),
        ],
        "preconditions": [
            "Written approval for production import from HOD and ICT.",
            "Official template downloaded from E-CIMES — column names unchanged.",
            "Backup or export snapshot taken before large imports.",
        ],
        "steps": [
            ("Data officer", "Download current template from Data → Import Data for the correct import type."),
            ("Data officer", "Populate spreadsheet from authoritative source; validate department, ward, and FY codes against catalogues."),
            ("HOD", "Sign import request with row count and effective date."),
            ("ICT Administrator", "Upload file for preview; resolve all validation errors — zero errors required."),
            ("ICT Administrator", "Confirm import; open Data Import Logs and save log reference."),
            ("M&E focal person", "Sample 5% of imported records in Projects Registry or Beneficiary Registry."),
            ("ICT Administrator", "Retain original file, corrected file, and log for audit."),
        ],
        "records": [
            "Import request approval.",
            "Preview error export (if any) and final clean file.",
            "Data Import Log entry with timestamp and user.",
        ],
        "controls": [
            "No production import from unverified spreadsheets.",
            "Imports blocked during critical reporting freeze unless exempted by Chief Officer.",
            "Duplicate detection review after project imports.",
        ],
        "escalation": "Failed or partial import → stop further imports → ICT + data owner root-cause meeting.",
        "manual_ref": "User Manual §8 Data import and upload logs.",
    },
    {
        "id": "SOP-012",
        "title": "ICT support, incidents, and system changes",
        "owner": "County ICT Administrator",
        "frequency": "Continuous; change windows scheduled",
        "purpose": (
            "Resolve E-CIMES incidents consistently, protect audit evidence, and control metadata "
            "and workflow changes that affect all users."
        ),
        "scope": "Help desk tickets, Audit Trail investigations, metadata/workflow changes, mobile app releases.",
        "roles": [
            ("End user", "Reports issue with required details per support checklist."),
            ("ICT help desk", "Triages, reproduces, resolves or escalates."),
            ("ICT Administrator", "Metadata, workflow, and release management."),
            ("Super Administrator", "Approves high-risk changes."),
        ],
        "preconditions": [
            "User completed first-line checks (login, filters, role, scope) per User Manual §3.",
            "Change request documented for metadata/workflow edits.",
        ],
        "steps": [
            ("End user", "Complete support checklist: name, department, role, module URL, error message, screenshot (masked), time."),
            ("ICT help desk", "Reproduce issue; check Audit Trail for related errors; assign priority (P1 system down → P4 cosmetic)."),
            ("ICT help desk", "Resolve or escalate to vendor/dev team with steps to reproduce."),
            ("ICT Administrator", "For metadata/workflow changes: test in non-production if available; announce maintenance window."),
            ("ICT Administrator", "Document change in change log; verify Audit Trail captures admin actions."),
            ("ICT Administrator", "Publish mobile APK releases via Admin; notify field teams to update."),
            ("ICT help desk", "Close ticket with resolution notes; review recurring issues monthly."),
        ],
        "records": [
            "Help desk ticket with resolution.",
            "Audit Trail extracts for investigations.",
            "Change log for metadata/workflow releases.",
        ],
        "controls": [
            "P1 incidents: response within 1 hour during business hours.",
            "No workflow changes on report deadline days without M&E + Planning approval.",
            "Quarterly review of AI Usage (Super Admin) for cost and policy compliance.",
        ],
        "escalation": "P1 unresolved 4 hours → ICT Manager → Chief Officer → vendor SLA escalation.",
        "manual_ref": "User Manual §15 Metadata and audit; §12 Support checklist.",
    },
]


def shade_cell(cell, fill: str) -> None:
    sh = OxmlElement("w:shd")
    sh.set(qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(sh)


def set_table_borders(table, color: str = "0D47A1") -> None:
    tbl_pr = table._tbl.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        table._tbl.insert(0, tbl_pr)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:color"), color)
        borders.append(el)
    tbl_pr.append(borders)


def set_bottom_border(paragraph, color: str = "1565C0", size: int = 8) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def setup_doc(doc: Document) -> None:
    for sec in doc.sections:
        sec.top_margin = Inches(0.85)
        sec.bottom_margin = Inches(0.75)
        sec.left_margin = Inches(0.9)
        sec.right_margin = Inches(0.9)
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)


def add_run_para(doc, text, *, bold=False, size=11, align=None, color=None, italic=False, space_after=6):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.name = "Calibri"
    r.font.size = Pt(size)
    if color:
        r.font.color.rgb = color
    return p


def section_heading(doc, text, level=1):
    if level == 1:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(8)
        r = p.add_run(text)
        r.bold = True
        r.font.size = Pt(16)
        r.font.color.rgb = NAVY
        set_bottom_border(p)
        return p
    p = doc.add_heading(text, level=level)
    for r in p.runs:
        r.font.color.rgb = NAVY
        r.font.size = Pt(13 if level == 2 else 12)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(item, style="List Bullet")
        for r in p.runs:
            r.font.name = "Calibri"
            r.font.size = Pt(10.5)


def add_labeled_block(doc, label, text, label_color):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    lr = p.add_run(f"{label}: ")
    lr.bold = True
    lr.font.color.rgb = label_color
    lr.font.size = Pt(10.5)
    tr = p.add_run(text)
    tr.font.size = Pt(10.5)


def add_table(doc, headers, rows, col_widths=None):
    if not rows:
        return
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    set_table_borders(table)
    for i, h in enumerate(headers):
        c = table.rows[0].cells[i]
        c.text = h
        shade_cell(c, HEADER_FILL)
        for p in c.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.name = "Calibri"
                r.font.size = Pt(9)
                r.font.color.rgb = RGBColor(255, 255, 255)
    for ri, row in enumerate(rows):
        fill = ALT_FILL if ri % 2 == 1 else None
        for ci, val in enumerate(row):
            c = table.rows[ri + 1].cells[ci]
            c.text = str(val)
            if fill:
                shade_cell(c, fill)
            for p in c.paragraphs:
                for r in p.runs:
                    r.font.name = "Calibri"
                    r.font.size = Pt(9)
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)
    doc.add_paragraph()


def cover_page(doc: Document) -> None:
    if LOGO_PATH.is_file():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(LOGO_PATH), width=Inches(1.15))

    add_run_para(doc, "REPUBLIC OF KENYA", bold=True, size=12, align=WD_ALIGN_PARAGRAPH.CENTER, color=SLATE)
    add_run_para(doc, "COUNTY GOVERNMENT OF MACHAKOS", bold=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER, color=NAVY)
    add_run_para(
        doc,
        "Department of Finance and Economic Planning",
        size=11,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        color=SLATE,
        space_after=18,
    )

    add_run_para(
        doc,
        "E-CIMES Standard Operating Procedures",
        bold=True,
        size=24,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        color=NAVY,
    )
    add_run_para(
        doc,
        "Electronic County Integrated Monitoring and Evaluation System",
        size=13,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        color=SLATE,
    )
    add_run_para(
        doc,
        "Official county processes for planning, projects, procurement,\n"
        "monitoring, finance, reporting, public engagement, and ICT support",
        size=11,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        italic=True,
        space_after=24,
    )

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for line in [
        f"Document version: {SOP_VERSION}",
        f"Effective date: {date.today().strftime('%d %B %Y')}",
        f"Document owner: {DOCUMENT_OWNER}",
        f"Review cycle: {REVIEW_CYCLE}",
        "Companion document: E-CIMES User Manual",
    ]:
        r = meta.add_run(line + "\n")
        r.font.size = Pt(10)
        r.font.color.rgb = SLATE

    doc.add_page_break()


def intro_section(doc: Document) -> None:
    section_heading(doc, "How this document relates to the User Manual")
    add_table(
        doc,
        ["Document", "Purpose", "Audience", "Answers the question"],
        [
            (
                "User Manual",
                "Explains system features, menus, and screens",
                "All staff (reference)",
                "How do I use this module?",
            ),
            (
                "Standard Operating Procedures (this document)",
                "Defines official county business processes using E-CIMES",
                "Process owners, HODs, auditors, ICT",
                "Who does what, when, with what approval and records?",
            ),
            (
                "User Training Script",
                "Voiceover script for video training",
                "Trainers, communications",
                "What do we say in the training video?",
            ),
        ],
        col_widths=[1.2, 1.6, 1.2, 1.2],
    )
    add_run_para(
        doc,
        "Staff should follow the SOPs for accountability and audit. Use the User Manual for step-by-step "
        "navigation when performing the tasks described in each SOP.",
        space_after=10,
    )

    section_heading(doc, "Document control")
    add_table(
        doc,
        ["Field", "Value"],
        [
            ("Title", "E-CIMES Standard Operating Procedures"),
            ("Version", SOP_VERSION),
            ("Effective date", date.today().strftime("%d %B %Y")),
            ("Owner", DOCUMENT_OWNER),
            ("Approved by", "[Chief Officer — Finance & Economic Planning]"),
            ("Review cycle", REVIEW_CYCLE),
            ("Distribution", "All departments, internal audit, county ICT, vendor support (reference)"),
        ],
        col_widths=[1.8, 4.4],
    )

    section_heading(doc, "SOP index")
    index_rows = [(s["id"], s["title"], s["owner"], s["frequency"]) for s in SOPS]
    add_table(doc, ["SOP ID", "Title", "Process owner", "Frequency"], index_rows, col_widths=[0.7, 2.5, 1.5, 1.5])
    doc.add_page_break()


def add_sop(doc: Document, sop: dict) -> None:
    section_heading(doc, f"{sop['id']}: {sop['title']}")

    add_table(
        doc,
        ["Control field", "Detail"],
        [
            ("SOP ID", sop["id"]),
            ("Title", sop["title"]),
            ("Process owner", sop["owner"]),
            ("Frequency", sop["frequency"]),
            ("Version", SOP_VERSION),
            ("Effective date", date.today().strftime("%d %B %Y")),
        ],
        col_widths=[1.5, 4.7],
    )

    section_heading(doc, "1.0 Purpose", level=2)
    add_run_para(doc, sop["purpose"], size=10.5)

    section_heading(doc, "2.0 Scope", level=2)
    add_run_para(doc, sop["scope"], size=10.5)

    section_heading(doc, "3.0 Roles and responsibilities", level=2)
    add_table(doc, ["Role", "Responsibility"], sop["roles"], col_widths=[1.6, 4.6])

    section_heading(doc, "4.0 Preconditions", level=2)
    add_bullets(doc, sop["preconditions"])

    section_heading(doc, "5.0 Procedure", level=2)
    proc_rows = [(str(i + 1), role, action) for i, (role, action) in enumerate(sop["steps"])]
    add_table(doc, ["Step", "Responsible role", "Action"], proc_rows, col_widths=[0.4, 1.4, 4.4])

    section_heading(doc, "6.0 Records and evidence", level=2)
    add_bullets(doc, sop["records"])

    section_heading(doc, "7.0 Quality controls", level=2)
    add_bullets(doc, sop["controls"])

    section_heading(doc, "8.0 Escalation", level=2)
    add_run_para(doc, sop["escalation"], size=10.5)

    section_heading(doc, "9.0 Related documents", level=2)
    add_bullets(doc, [sop["manual_ref"], "E-CIMES User Manual (docs/E-CIMES-User-Manual.docx)", "County records management policy"])

    doc.add_page_break()


def revision_history(doc: Document) -> None:
    section_heading(doc, "Revision history")
    add_table(
        doc,
        ["Version", "Date", "Author", "Summary of changes", "Approved by"],
        [
            (SOP_VERSION, date.today().strftime("%d %B %Y"), "County ICT / M&E", "Initial release", "[Pending signature]"),
        ],
        col_widths=[0.6, 0.9, 1.2, 2.0, 1.5],
    )
    add_run_para(
        doc,
        "— End of Standard Operating Procedures —",
        bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        color=NAVY,
        space_after=6,
    )
    add_run_para(
        doc,
        "County Government of Machakos · E-CIMES\n"
        "For system navigation detail, refer to the E-CIMES User Manual.",
        size=9,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        italic=True,
        color=SLATE,
    )


def build_sop_manual() -> Document:
    doc = Document()
    setup_doc(doc)
    cover_page(doc)
    intro_section(doc)
    for sop in SOPS:
        add_sop(doc, sop)
    revision_history(doc)
    return doc


def export_pdf(docx_path: Path, pdf_path: Path) -> bool:
    soffice = shutil.which("libreoffice") or shutil.which("soffice")
    if not soffice:
        return False
    try:
        subprocess.run(
            [soffice, "--headless", "--convert-to", "pdf", str(docx_path), "--outdir", str(pdf_path.parent)],
            check=True,
            capture_output=True,
            timeout=120,
        )
    except (subprocess.CalledProcessError, subprocess.TimeoutExpired):
        return False
    return pdf_path.is_file()


def main():
    OUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc = build_sop_manual()
    doc.save(OUT_DOCX)
    print(f"Wrote {OUT_DOCX}")
    print(f"  SOPs: {len(SOPS)}")
    if export_pdf(OUT_DOCX, OUT_PDF):
        print(f"Wrote {OUT_PDF}")
    else:
        print("  PDF: install LibreOffice or open the .docx → File → Export as PDF.")


if __name__ == "__main__":
    main()
