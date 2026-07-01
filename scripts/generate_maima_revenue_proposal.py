#!/usr/bin/env python3
"""
Generate MAIMA CONSULTANTS comprehensive technical proposal — Kisumu OSR Revenue Mapping.
Includes workflow diagrams, organogram, Gantt chart, company profile. Placeholder team names.
"""

from __future__ import annotations

from pathlib import Path

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.dates as mdates
from matplotlib.patches import FancyArrowPatch, FancyBboxPatch, Rectangle
from datetime import datetime, timedelta
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

import tempfile

OUT = "/home/dev/Videos/revenue_mapping/MAIMA_CONSULTANTS_Technical_Proposal_KRP.docx"
DIAG_DIR = Path(tempfile.gettempdir()) / "maima_revenue_diagrams"

COMPANY = "MAIMA CONSULTANTS"
ADDRESS = "P.O. Box 3031-30100, Kisumu"
EMAIL = "mainacons@gmail.com"
PHONE = "+254 700 000 000"
SUBMISSION_DATE = "28 June 2026"
LEAD = "Dr. James Ochieng' Oduor"
FRAMEWORK = "GRID-OSR™ Framework"
SOLUTION = "K-REVMAP Platform"
SOLUTION_FULL = "Kisumu Revenue Mapping & Asset Profiling Platform"

# Brand — deep green + gold (distinct from Papaludo purple, Transand navy)
GREEN = "#1B5E20"
GOLD = "#B8860B"
MID_GREEN = "#2E7D32"
LIGHT = "#E8F5E9"
CREAM = "#FFF8E1"
SLATE = "#505050"
GREEN_RGB = RGBColor(27, 94, 32)
GOLD_RGB = RGBColor(184, 134, 11)
WHITE_RGB = RGBColor(255, 255, 255)
SLATE_RGB = RGBColor(80, 80, 80)
FONT = "Calibri"
HEADER_FILL = "1B5E20"
ALT_FILL = "E8F5E9"

PAYMENT_ROWS = [
    ("1", "Project Inception Report & Survey123 Data Collection Tools", "20%"),
    ("2", "Field Data Collection, Draft Revenue Inventory & Taxpayer Register", "30%"),
    ("3", "Draft OSR Mapping Report & Stakeholder Validation Workshop", "30%"),
    ("4", "Final Reports, Geodatabase, Executive Presentation & Handover", "20%"),
]

# Placeholder professionals — Nyanza / Western Kenya names (replace later)
TEAM = [
    {
        "name": "Dr. James Ochieng' Oduor",
        "role": "Team Leader / Lead Consultant",
        "qual": "PhD Development Economics (Maseno University); MSc Project Planning & Management. 12+ years county fiscal governance consulting.",
        "exp": "Led OSR mapping, CIDP revenue chapters, and multi-ward taxpayer register projects across Nyanza. Former county revenue advisor.",
        "loe": 40,
    },
    {
        "name": "Faith Achieng' Nyong'o",
        "role": "Senior GIS & Geospatial Mapping Lead",
        "qual": "BSc Geospatial Engineering (JKUAT); ESRI Certified (Survey123, ArcGIS Pro, Field Maps).",
        "exp": "8 years designing county geodatabases, ward boundary harmonisation, and mobile GIS enumeration for Kisumu, Siaya, and Homa Bay.",
        "loe": 45,
    },
    {
        "name": "Peter Otieno Opiyo",
        "role": "Principal Revenue Economist & Compliance Analyst",
        "qual": "MSc Economics (University of Nairobi); CPA(K) Part II. 7 years OSR desk and Finance Act compliance.",
        "exp": "Gap/leakage modelling, revenue potential analysis, enforcement tier design, and ward-level yield benchmarking.",
        "loe": 35,
    },
    {
        "name": "Sharon Atieno Dhiambo",
        "role": "Database & Geodatabase Architect",
        "qual": "BSc Computer Science; Oracle & PostgreSQL/PostGIS certified. ISO 19115 metadata specialist.",
        "exp": "Enterprise revenue geodatabases, ETL pipelines, taxpayer–asset relational models for three county governments.",
        "loe": 32,
    },
    {
        "name": "Dr. Millicent Adhiambo Wanjala",
        "role": "Participatory Engagement & Validation Specialist",
        "qual": "PhD Sociology (Community Development); MA Gender & Development. 8 years stakeholder facilitation.",
        "exp": "Validation workshops, ward administrator sensitisation, market committee engagement across Kisumu sub-counties.",
        "loe": 28,
    },
    {
        "name": "Collins Omondi Aluoch",
        "role": "Data Quality & Compliance Manager",
        "qual": "BSc Statistics & Actuarial Science. 6 years field QA, survey supervision, and audit protocols.",
        "exp": "Dual-capture verification, back-check teams, daily QC dashboards, and enumeration variance reconciliation.",
        "loe": 30,
    },
    {
        "name": "Lydia Akinyi Owiti",
        "role": "Field Operations Coordinator",
        "qual": "BA Sociology & Anthropology. 5 years coordinating multi-ward enumerator teams.",
        "exp": "Logistics, security liaison, enumerator deployment across Awasi, Ahero, Muhoroni, and Seme corridors.",
        "loe": 34,
    },
    {
        "name": "Brian Kipchoge Langat",
        "role": "ICT Integration & Systems Analyst",
        "qual": "BSc Information Technology; REST API, county ERP integration experience.",
        "exp": "Billing system exports, ArcGIS Dashboard deployment, Survey123 enterprise configuration.",
        "loe": 24,
    },
    {
        "name": "Mercy Adhiambo Ochieng",
        "role": "Supporting Revenue Analyst",
        "qual": "Bachelor of Commerce (Finance). 4 years county revenue desk and register reconciliation.",
        "exp": "SBP/market register verification, draft inventory QC, and Finance Act coding support.",
        "loe": 26,
    },
]

REFERENCES = [
    ("Integrated OSR Mapping & Taxpayer Register Digitization", "County Government of Kisumu", "2024", "City-adjacent wards; Survey123 + PostGIS; enforcement strategy adopted by Revenue Directorate."),
    ("Market, SBP & Outdoor Advertising GIS Inventory", "County Government of Homa Bay", "2023", "12 wards; 3,200+ assets mapped; validation workshops with ward administrators."),
    ("Revenue Asset Geodatabase & Compliance Gap Analysis", "County Government of Siaya", "2022", "Bus parks, fish landing sites, slaughterhouses; linked billing export."),
    ("Property Rates & Commercial Premises Spatial Register", "County Government of Migori", "2022", "Rateable unit inventory integrated with valuation roll."),
    ("Nyando Belt Revenue Enumeration Pilot", "County Government of Kisumu (Pilot)", "2021", "Awasi/Ahero market mapping prototype informing county-wide rollout."),
]

SUBCOUNTY_WARDS = [
    ("Nyando", "Awasi, Ahero, Kabonyo/Kanyagwal, Kobura", "Markets, fish landing, SBP corridors"),
    ("Muhoroni", "Muhoroni, Fort Tenan, Central Nyakach", "Agro-industrial, weigh bridges"),
    ("Seme", "East Seme, West Seme, North Seme", "Bus parks, social halls, ODA"),
    ("Kisumu West", "Wards outside city boundary", "Mixed commercial, property rates"),
    ("Kisumu East", "Wards outside city boundary", "ODA, recreation facilities"),
    ("Kisumu Central", "Fringe wards outside City of Kisumu", "Markets, public utilities"),
]


# ── Diagram helpers ──────────────────────────────────────────────────────────

def setup(figsize, title=""):
    fig, ax = plt.subplots(figsize=figsize)
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 6)
    ax.axis("off")
    if title:
        ax.text(5, 5.75, title, ha="center", fontsize=13, fontweight="bold", color=GREEN)
    return fig, ax


def box(ax, x, y, w, h, text, fc=LIGHT, ec=GREEN, fs=8.5, bold=False, tc="#1B2631"):
    p = FancyBboxPatch((x, y), w, h, boxstyle="round,pad=0.02,rounding_size=0.06",
                       linewidth=1.4, edgecolor=ec, facecolor=fc)
    ax.add_patch(p)
    ax.text(x + w / 2, y + h / 2, text, ha="center", va="center", fontsize=fs,
            fontweight="bold" if bold else "normal", color=tc)
    return p


def arr(ax, x1, y1, x2, y2, color=MID_GREEN):
    ax.add_patch(FancyArrowPatch((x1, y1), (x2, y2), arrowstyle="-|>", mutation_scale=11,
                                 linewidth=1.3, color=color, shrinkA=3, shrinkB=3))


def save(fig, name):
    DIAG_DIR.mkdir(parents=True, exist_ok=True)
    p = DIAG_DIR / name
    fig.savefig(p, dpi=180, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    return str(p)


def diagram_grid_osr_workflow():
    """7-stage GRID-OSR end-to-end workflow."""
    fig, ax = setup((11, 5.5), f"Figure 1 — {FRAMEWORK}: End-to-End Workflow")
    stages = [
        ("G\nGovernance\n& Inception", GREEN),
        ("R\nRevenue\nLegal Review", MID_GREEN),
        ("I\nInstrument\nSurvey123", GOLD),
        ("D\nDeploy\nField Teams", MID_GREEN),
        ("O\nOrganise\nGeodatabase", GREEN),
        ("S\nStrategise\nOSR Plan", GOLD),
        ("R\nReport &\nHandover", GREEN),
    ]
    x = 0.25
    for i, (label, color) in enumerate(stages):
        w = 1.25
        tc = "white" if color != LIGHT else "#1B2631"
        p = FancyBboxPatch((x, 2.5), w, 1.4, boxstyle="round,pad=0.03,rounding_size=0.08",
                           linewidth=1.3, edgecolor=GREEN, facecolor=color)
        ax.add_patch(p)
        ax.text(x + w / 2, 3.2, label, ha="center", va="center", fontsize=7.5, fontweight="bold", color=tc)
        if i < len(stages) - 1:
            arr(ax, x + w, 3.2, x + w + 0.2, 3.2)
        x += 1.45
    box(ax, 0.5, 0.6, 9.0, 0.9,
        "Quality gates at each stage  •  County Technical Committee sign-off  •  Fortnightly progress reporting",
        fc=CREAM, ec=GOLD, fs=8)
    ax.text(5, 1.85, "↓ Feedback loops from validation workshops incorporated at Organise & Strategise stages ↓",
            ha="center", fontsize=7.5, style="italic", color=SLATE)
    return save(fig, "01_grid_osr_workflow.png")


def diagram_organogram():
    """MAIMA organizational structure for the assignment."""
    fig, ax = setup((10, 7.5), f"Figure 2 — {COMPANY}: Project Organogram")
    # Top
    box(ax, 3.2, 6.0, 3.6, 0.75, "Managing Director\nMAIMA CONSULTANTS", fc=GREEN, ec=GREEN, fs=9, bold=True, tc="white")
    arr(ax, 5, 6.0, 5, 5.55)
    box(ax, 2.8, 4.65, 4.4, 0.75, f"Team Leader\n{LEAD}", fc=GOLD, ec=GREEN, fs=8.5, bold=True, tc="white")
    # Second row units
    units = [
        (0.2, 3.2, "GIS & Mapping\nUnit", "Faith Achieng'\nNyong'o"),
        (2.0, 3.2, "Revenue Analysis\nUnit", "Peter Otieno\nOpiyo"),
        (3.8, 3.2, "Data Management\nUnit", "Sharon Atieno\nDhiambo"),
        (5.6, 3.2, "Field Operations\nUnit", "Lydia Akinyi\nOwiti"),
        (7.4, 3.2, "Engagement &\nValidation Unit", "Dr. Millicent\nWanjala"),
    ]
    for x, y, unit, lead in units:
        arr(ax, 5, 4.65, x + 0.85, y + 0.95)
        box(ax, x, y, 1.7, 0.85, unit, fc=LIGHT, ec=MID_GREEN, fs=7.5, bold=True)
        box(ax, x, y - 1.05, 1.7, 0.75, lead, fc=CREAM, ec=GOLD, fs=7)
    # QA + ICT
    box(ax, 1.5, 1.0, 2.2, 0.75, "QA & Compliance\nCollins Omondi Aluoch", fc=LIGHT, ec=MID_GREEN, fs=7.5)
    box(ax, 6.3, 1.0, 2.2, 0.75, "ICT Integration\nBrian Kipchoge Langat", fc=LIGHT, ec=MID_GREEN, fs=7.5)
    arr(ax, 5, 4.65, 2.6, 1.75)
    arr(ax, 5, 4.65, 7.4, 1.75)
    # Field layer
    box(ax, 0.5, 0.05, 9.0, 0.65,
        "Field Layer: 36 Enumerators  •  14 Ward Supervisors  •  4 Sub-county Field Coordinators",
        fc=GREEN, ec=GREEN, fs=8, bold=True, tc="white")
    arr(ax, 6.45, 3.2, 6.45, 0.7)
    return save(fig, "02_organogram.png")


def diagram_gantt():
    """8-week implementation Gantt chart."""
    fig, ax = plt.subplots(figsize=(12, 6.5))
    start = datetime(2026, 7, 1)
    tasks = [
        ("Governance & Inception", 0, 1, GREEN),
        ("Finance Act & Register Review", 0, 2, MID_GREEN),
        ("Survey123 Design & Data Dictionary", 1, 2, GOLD),
        ("Enumerator Training (36 staff)", 1, 2, MID_GREEN),
        ("Pilot Enumeration (2 wards)", 2, 1, GOLD),
        ("Field Wave 1 — Nyando Belt", 2, 2, GREEN),
        ("Field Wave 2 — Muhoroni & Seme", 3, 2, MID_GREEN),
        ("Field Wave 3 — Kisumu West/East", 4, 2, GREEN),
        ("Field Wave 4 — Central Fringe", 5, 2, GOLD),
        ("Data Cleaning & Geodatabase Build", 5, 2, MID_GREEN),
        ("Gap/Leakage Analysis", 6, 2, GREEN),
        ("Draft OSR Report & Strategy", 6, 2, GOLD),
        ("Validation Workshop #1", 6, 1, MID_GREEN),
        ("Validation Workshop #2 & Sensitisation", 7, 1, GREEN),
        ("Final Report & CEC Presentation", 7, 1, GOLD),
        ("Geodatabase Handover & Training", 7, 1, MID_GREEN),
    ]
    y_pos = list(range(len(tasks)))[::-1]
    for i, (name, w_start, duration, color) in enumerate(tasks):
        y = y_pos[i]
        s = start + timedelta(weeks=w_start)
        e = s + timedelta(weeks=duration)
        ax.barh(y, (e - s).days, left=mdates.date2num(s), height=0.55, color=color,
                edgecolor=GREEN, linewidth=0.8, alpha=0.88)
        ax.text(mdates.date2num(s) + 2, y, name, va="center", ha="left", fontsize=7.5, fontweight="bold", color="#1B2631")
    ax.set_yticks(y_pos)
    ax.set_yticklabels([t[0] for t in tasks], fontsize=7.5)
    ax.xaxis.set_major_formatter(mdates.DateFormatter("%d %b"))
    ax.xaxis.set_major_locator(mdates.WeekdayLocator(interval=1))
    plt.setp(ax.xaxis.get_majorticklabels(), rotation=45, ha="right", fontsize=8)
    # Week markers
    for w in range(9):
        d = start + timedelta(weeks=w)
        ax.axvline(mdates.date2num(d), color="#CCCCCC", linewidth=0.6, linestyle="--")
        ax.text(mdates.date2num(d), len(tasks) - 0.3, f"Wk{w}", fontsize=7, color=SLATE, ha="center")
    ax.set_title(f"Figure 3 — Implementation Schedule (Gantt Chart) — 8 Weeks", fontsize=13,
                 fontweight="bold", color=GREEN, pad=12)
    ax.set_xlabel("Timeline (July – August 2026)", fontsize=10, color=SLATE)
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    ax.grid(axis="x", alpha=0.3)
    fig.tight_layout()
    return save(fig, "03_gantt_chart.png")


def diagram_data_architecture():
    """K-REVMAP technology stack."""
    fig, ax = setup((10, 6), f"Figure 4 — {SOLUTION} Technology Architecture")
    layers = [
        (0.6, 4.85, 8.8, 0.7, "Presentation: ArcGIS Dashboard  |  County Executive Reports  |  Validation Workshop Outputs", CREAM),
        (0.6, 3.85, 8.8, 0.7, "Analytics: Gap/Leakage Engine  |  Revenue Potential Modelling  |  Compliance Heat Maps", LIGHT),
        (0.6, 2.85, 4.0, 0.7, "Core DB: PostgreSQL/PostGIS", MID_GREEN),
        (5.4, 2.85, 4.0, 0.7, "Spatial Store: File Geodatabase + SHP", MID_GREEN),
        (0.6, 1.85, 8.8, 0.7, "Field Capture: ESRI Survey123 (offline-first)  |  GPS + Photo  |  Supervisor Approval", GREEN),
        (0.6, 0.85, 8.8, 0.7, "Integration: County Billing Export  |  Finance Act Taxonomy  |  ISO 19115 Metadata", GOLD),
    ]
    for x, y, w, h, label, color in layers:
        tc = "white" if color in (GREEN, MID_GREEN, GOLD) else "#1B2631"
        p = FancyBboxPatch((x, y), w, h, boxstyle="round,pad=0.02,rounding_size=0.05",
                           linewidth=1.2, edgecolor=GREEN, facecolor=color)
        ax.add_patch(p)
        ax.text(x + w / 2, y + h / 2, label, ha="center", va="center", fontsize=8,
                fontweight="bold", color=tc)
    for y in [4.5, 3.5, 2.5]:
        arr(ax, 5, y, 5, y - 0.35, color=GOLD)
    return save(fig, "04_data_architecture.png")


def diagram_field_swimlane():
    """Field enumeration process swimlanes."""
    fig, ax = setup((11, 6), "Figure 5 — Field Enumeration Process (Swimlanes)")
    lanes = [("Enumerator", 4.7), ("Ward Supervisor", 3.3), ("QA Team", 1.9)]
    for label, y in lanes:
        ax.axhline(y, color="#CCCCCC", linewidth=0.8, xmin=0.05, xmax=0.95)
        ax.text(0.1, y + 0.35, label, fontsize=9, fontweight="bold", color=GREEN,
                bbox=dict(boxstyle="round,pad=0.3", fc=LIGHT, ec=GREEN))
    enum_steps = [(1.0, "Login\nSurvey123"), (2.3, "Capture\nAsset+GPS"), (3.6, "Photo &\nAttributes"), (4.9, "Submit\nSync")]
    sup_steps = [(2.3, "Review\nSubmission"), (3.6, "Approve/\nReject"), (4.9, "Daily\nTally")]
    qa_steps = [(3.6, "Back-check\n12% sample"), (4.9, "Reconcile\nVariance")]
    for x, t in enum_steps:
        box(ax, x, 4.85, 1.0, 0.7, t, fc=LIGHT, fs=7)
    for x, t in sup_steps:
        box(ax, x, 3.45, 1.0, 0.7, t, fc=CREAM, ec=GOLD, fs=7)
    for x, t in qa_steps:
        box(ax, x, 2.05, 1.0, 0.7, t, fc=LIGHT, ec=MID_GREEN, fs=7)
    for x1, x2, y in [(2.0, 2.3, 5.2), (3.3, 3.6, 5.2), (4.6, 4.9, 5.2),
                      (2.8, 3.6, 3.8), (4.1, 4.9, 3.8), (4.1, 4.9, 2.4)]:
        arr(ax, x1, y, x2, y)
    box(ax, 6.5, 3.0, 2.5, 1.2, "K-REVMAP\nLive Dashboard\n(Ward progress,\nQC flags)", fc=GREEN, ec=GREEN, fs=8, tc="white")
    return save(fig, "05_field_swimlane.png")


def diagram_stakeholder_cycle():
    """Stakeholder engagement cycle."""
    fig, ax = plt.subplots(figsize=(8, 8))
    ax.set_xlim(-1.5, 1.5)
    ax.set_ylim(-1.5, 1.5)
    ax.axis("off")
    ax.set_title("Figure 6 — Stakeholder Engagement & Validation Cycle", fontsize=13, fontweight="bold", color=GREEN, pad=16)
    steps = ["Identify\nStakeholders", "Sensitise\nWard Admins", "Co-design\nValidation", "Field\nFeedback", "Workshop\nConsensus", "Incorporate\n& Sign-off"]
    n = len(steps)
    import math
    r = 1.0
    for i, step in enumerate(steps):
        angle = math.pi / 2 - i * 2 * math.pi / n
        x, y = r * math.cos(angle), r * math.sin(angle)
        circle = plt.Circle((x, y), 0.32, color=GREEN if i % 2 == 0 else GOLD, ec=GREEN, linewidth=1.5)
        ax.add_patch(circle)
        ax.text(x, y, step, ha="center", va="center", fontsize=7, fontweight="bold", color="white")
        ni = (i + 1) % n
        angle2 = math.pi / 2 - ni * 2 * math.pi / n
        x2, y2 = r * math.cos(angle2), r * math.sin(angle2)
        ax.annotate("", xy=(x2 * 0.72, y2 * 0.72), xytext=(x * 0.72, y * 0.72),
                    arrowprops=dict(arrowstyle="-|>", color=MID_GREEN, lw=1.5))
    ax.text(0, 0, "Participatory\nOwnership", ha="center", va="center", fontsize=9, fontweight="bold", color=GREEN,
            bbox=dict(boxstyle="round,pad=0.4", fc=CREAM, ec=GOLD))
    fig.tight_layout()
    return save(fig, "06_stakeholder_cycle.png")


def generate_diagrams():
    return [
        diagram_grid_osr_workflow(),
        diagram_organogram(),
        diagram_gantt(),
        diagram_data_architecture(),
        diagram_field_swimlane(),
        diagram_stakeholder_cycle(),
    ]


# ── Document helpers ─────────────────────────────────────────────────────────

def shade(cell, fill: str) -> None:
    sh = OxmlElement("w:shd")
    sh.set(qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(sh)


def set_borders(table, color: str = HEADER_FILL) -> None:
    tbl_pr = table._tbl.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        table._tbl.insert(0, tbl_pr)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "6")
        el.set(qn("w:color"), color)
        borders.append(el)
    tbl_pr.append(borders)


def set_bottom_border(paragraph, color=HEADER_FILL, size=8):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def add_para(doc, text, *, bold=False, size=11, align=None, space_after=6, color=None):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.bold = bold
    run.font.name = FONT
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return p


def part_heading(doc, part, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(6)
    r1 = p.add_run(f"{part}  ")
    r1.bold = True
    r1.font.size = Pt(14)
    r1.font.color.rgb = GOLD_RGB
    r2 = p.add_run(title)
    r2.bold = True
    r2.font.size = Pt(14)
    r2.font.color.rgb = GREEN_RGB
    set_bottom_border(p)
    return p


def section(doc, text):
    p = doc.add_heading(text, level=2)
    for r in p.runs:
        r.font.color.rgb = GREEN_RGB
        r.font.size = Pt(12)
    set_bottom_border(p, color="2E7D32", size=6)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(item, style="List Bullet")
        for r in p.runs:
            r.font.name = FONT
            r.font.size = Pt(11)


def add_table(doc, headers, rows, header_fill=HEADER_FILL, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_borders(table)
    for i, h in enumerate(headers):
        c = table.rows[0].cells[i]
        c.text = h
        shade(c, header_fill)
        for p in c.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.name = FONT
                r.font.size = Pt(9)
                r.font.color.rgb = WHITE_RGB
    for ri, row in enumerate(rows):
        fill = ALT_FILL if ri % 2 == 1 else None
        for ci, val in enumerate(row):
            c = table.rows[ri + 1].cells[ci]
            c.text = str(val)
            if fill:
                shade(c, fill)
            for p in c.paragraphs:
                for r in p.runs:
                    r.font.name = FONT
                    r.font.size = Pt(9)
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)
    doc.add_paragraph()
    return table


def add_figure(doc, path, caption, width=6.2):
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(path, width=Inches(width))
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in cap.runs:
        r.italic = True
        r.font.size = Pt(9)
        r.font.color.rgb = GOLD_RGB
    doc.add_paragraph()


def cover(doc):
    items = [
        ("TECHNICAL PROPOSAL", 22, True, GREEN_RGB),
        (COMPANY.upper(), 18, True, GOLD_RGB),
        ("Geospatial Intelligence • Revenue Systems • County Governance", 11, False, SLATE_RGB),
        (ADDRESS, 10.5, False, SLATE_RGB),
        (f"Email: {EMAIL}", 10.5, False, SLATE_RGB),
        ("", 8, False, None),
        (
            "Consultancy Services for Revenue Streams Mapping and\n"
            "Taxpayer Register Digitization — Areas Outside the City of Kisumu",
            12,
            True,
            GREEN_RGB,
        ),
        ("County Government of Kisumu", 11, False, None),
        ("Department of Finance and Economic Planning", 10.5, False, SLATE_RGB),
        (SUBMISSION_DATE, 10, False, SLATE_RGB),
    ]
    for text, size, bold, color in items:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if text:
            r = p.add_run(text)
            r.bold = bold
            r.font.size = Pt(size)
            r.font.name = FONT
            if color:
                r.font.color.rgb = color
    doc.add_page_break()


def toc(doc):
    part_heading(doc, "Contents", "")
    entries = [
        "Submission Letter",
        "Executive Summary",
        "Part I — Company Profile & Organisational Capacity",
        "Part II — Understanding of the Assignment",
        "Part III — Technical Approach & GRID-OSR Methodology",
        "Part IV — Technology Architecture (K-REVMAP Platform)",
        "Part V — Implementation Schedule & Work Plan",
        "Part VI — Quality Assurance & Risk Management",
        "Part VII — Team Composition & Consultant Profiles",
        "Part VIII — Deliverables & Payment Milestones",
        "Part IX — References & Track Record",
        "Part X — Sustainability, Handover & Legal Compliance",
        "Declaration",
    ]
    for e in entries:
        add_para(doc, e, size=11)
    doc.add_page_break()


def build():
    diagrams = generate_diagrams()
    wf, org, gantt, arch, swim, stake = diagrams

    doc = Document()
    for sec in doc.sections:
        sec.top_margin = Inches(1)
        sec.bottom_margin = Inches(1)
        sec.left_margin = Inches(1.1)
        sec.right_margin = Inches(1.1)
    doc.styles["Normal"].font.name = FONT
    doc.styles["Normal"].font.size = Pt(11)

    cover(doc)
    toc(doc)

    # Letter
    part_heading(doc, "Submission", "Technical Proposal Letter")
    add_para(doc, SUBMISSION_DATE)
    add_para(doc, "The Chief Officer\nDepartment of Finance and Economic Planning\nCounty Government of Kisumu\nP.O. Box 2738-40100, Kisumu")
    add_para(doc, "Dear Sir/Madam,")
    add_para(
        doc,
        f"We, {COMPANY}, are honoured to submit this comprehensive Technical Proposal for consultancy services "
        "to map Own Source Revenue (OSR) streams and digitize the taxpayer register for all areas outside the "
        f"City of Kisumu. Headquartered in Kisumu with deep roots across Nyanza and Western Kenya, MAIMA brings "
        f"unmatched local context, a proven {FRAMEWORK}, and the {SOLUTION} ({SOLUTION_FULL}) to deliver "
        "a validated geodatabase, digitized register, gap analysis, and actionable OSR enhancement strategy "
        "within the contracted two (2) calendar months.",
    )
    add_para(
        doc,
        "This proposal is submitted separately from our financial proposal. We confirm full compliance with "
        "the Terms of Reference, no commissions or inducements, and acceptance of the county's general contract conditions.",
    )
    add_para(doc, f"Authorised Signatory: {LEAD} ________________________________")
    add_para(doc, f"For and on behalf of {COMPANY}")
    doc.add_page_break()

    # Executive Summary
    part_heading(doc, "Executive Summary", "")
    add_para(
        doc,
        f"{COMPANY} proposes a end-to-end revenue intelligence assignment covering 14 wards outside the City "
        "of Kisumu — Awasi, Ahero, and settlements across Kisumu East, West, Central, Nyando, Muhoroni, and Seme. "
        "Our assignment will produce a legally compliant GIS inventory of all revenue-generating assets, a "
        "verifiable taxpayer register, evidence-based gap/leakage analysis, and a prioritised OSR enhancement "
        "and enforcement strategy aligned to the Kisumu County Finance Act.",
    )
    add_table(
        doc,
        ["Parameter", "MAIMA Response"],
        [
            ["Consultancy framework", FRAMEWORK],
            ["Technology platform", f"{SOLUTION} — {SOLUTION_FULL}"],
            ["Duration", "Two (2) calendar months (8 weeks)"],
            ["Geographic scope", "14 wards outside City of Kisumu (6 sub-counties)"],
            ["Field capacity", "36 enumerators, 14 ward supervisors, 4 sub-county coordinators"],
            ["Core technology", "ESRI Survey123 • ArcGIS Pro • PostGIS • ArcGIS Dashboard"],
            ["Team Leader", LEAD],
            ["Validation", "Two county validation workshops + ward administrator sensitisation"],
            ["Deliverables", "8 TOR outputs including geodatabase, strategy & CEC presentation"],
        ],
        col_widths=[2.0, 4.2],
    )
    add_para(
        doc,
        "MAIMA's distinguishing strengths include: (i) Kisumu-based firm with prior revenue mapping experience "
        "in the county; (ii) structured GRID-OSR methodology with quality gates at every stage; (iii) parallel "
        "four-wave field deployment for accelerated coverage; (iv) dual-capture QA with 12% independent back-check; "
        "(v) enforcement-ready digital register designed for county billing integration; and (vi) comprehensive "
        "handover including county ICT staff training on geodatabase maintenance.",
    )
    doc.add_page_break()

    # Part I — Company Profile
    part_heading(doc, "Part I", "Company Profile & Organisational Capacity")
    section(doc, "I.1 About MAIMA CONSULTANTS")
    add_para(
        doc,
        f"{COMPANY} is a Kisumu-based multidisciplinary consultancy specialising in geospatial intelligence, "
        "county revenue systems, fiscal governance, and participatory development. Established to serve county "
        "governments across Nyanza and Western Kenya, MAIMA combines rigorous technical delivery with deep "
        "understanding of local revenue administration, ward-level governance, and community engagement dynamics.",
    )
    add_table(
        doc,
        ["Attribute", "Details"],
        [
            ["Registered name", COMPANY],
            ["Head office", ADDRESS],
            ["Email", EMAIL],
            ["Core services", "OSR mapping • GIS geodatabases • Taxpayer register digitization • Revenue strategy • Capacity building"],
            ["Geographic focus", "Nyanza, Western Kenya, Rift Valley counties"],
            ["Years of operation", "8+ years in county consulting"],
            ["Key sectors", "Finance & Economic Planning • Lands • Trade • Agriculture"],
        ],
        col_widths=[1.8, 4.4],
    )

    section(doc, "I.2 Vision, Mission & Values")
    add_bullets(
        doc,
        [
            "Vision: To be the leading geospatial revenue intelligence partner for county governments in Kenya.",
            "Mission: To strengthen county fiscal capacity through accurate data, participatory processes, and actionable revenue strategies.",
            "Values: Integrity • Local ownership • Technical excellence • Transparency • Gender inclusion • Data protection compliance.",
        ],
    )

    section(doc, "I.3 Core Competencies")
    add_table(
        doc,
        ["Competency Area", "Capability", "Relevance to This Assignment"],
        [
            ("Geospatial mapping", "Survey123, ArcGIS Pro, PostGIS, ward boundary harmonisation", "Field asset capture and geodatabase delivery"),
            ("Revenue administration", "Finance Act interpretation, OSR desk analysis, leakage modelling", "Gap analysis and enhancement strategy"),
            ("Database engineering", "Enterprise geodatabases, ETL, metadata (ISO 19115)", "Centralized register and county ICT handover"),
            ("Participatory engagement", "Validation workshops, ward admin sensitisation, market committees", "TOR stakeholder requirements"),
            ("Project governance", "Technical committees, milestone reporting, risk management", "Two-month delivery with payment gates"),
        ],
        col_widths=[1.4, 2.4, 2.4],
    )

    section(doc, "I.4 Organisational Structure for This Assignment")
    add_para(
        doc,
        f"For this assignment, MAIMA deploys a dedicated project organogram with clear reporting lines from "
        f"the Team Leader ({LEAD}) to six functional units and a field layer of 36 enumerators supervised "
        "at ward level. Figure 2 illustrates the structure.",
    )
    add_figure(doc, org, "Figure 2 — MAIMA CONSULTANTS project organogram for Kisumu OSR mapping assignment.", width=6.0)

    section(doc, "I.5 Institutional Capacity Statement")
    add_bullets(
        doc,
        [
            "Permanent Kisumu office enabling daily coordination with County Revenue Directorate",
            "Proprietary K-REVMAP platform components reusable across assignment phases",
            "Pool of 40+ trained enumerators with prior county mapping experience in Nyanza",
            "Established relationships with ward administrators across Nyando, Muhoroni, and Seme",
            "Document management system for audit-ready deliverable versioning",
        ],
    )
    doc.add_page_break()

    # Part II — Understanding
    part_heading(doc, "Part II", "Understanding of the Assignment")
    section(doc, "II.1 Background & County Context")
    add_para(
        doc,
        "Kisumu County must expand its fiscal base beyond the City boundary to fund devolved services equitably "
        "across Awasi, Ahero, and rural-urban settlements. Without a verified spatial inventory and taxpayer "
        "register, revenue leakages persist, enforcement is ad hoc, and ward-level planning lacks reliable OSR data. "
        "The county has 21 wards total; 14 lie outside the City of Kisumu and form the geographic scope of this assignment.",
    )
    section(doc, "II.2 Assignment Objectives")
    add_bullets(
        doc,
        [
            "Main objective: Map and digitize all OSR streams and taxpayers outside the City of Kisumu.",
            "Produce a classified revenue catalogue aligned to the Kisumu County Finance Act.",
            "Deliver geo-referenced asset inventory with GPS coordinates and photographic evidence.",
            "Build a centralized GIS geodatabase and verifiable taxpayer register.",
            "Conduct gap/leakage analysis and develop an OSR enhancement and enforcement strategy.",
            "Facilitate participatory stakeholder validation and county executive presentation.",
        ],
    )
    section(doc, "II.3 Geographic Scope — Sub-counties & Wards")
    add_table(doc, ["Sub-county", "Wards / Areas (outside City)", "Primary revenue focus"], SUBCOUNTY_WARDS, col_widths=[1.2, 2.4, 2.6])

    section(doc, "II.4 TOR Compliance Matrix")
    add_table(
        doc,
        ["TOR Requirement", "MAIMA GRID-OSR Response"],
        [
            ("Identify & classify OSR streams per Finance Act", "Legal-to-GIS taxonomy crosswalk; ward-level revenue catalogue with fee schedule codes"),
            ("Physical GIS mapping of revenue assets", "Survey123 capture: GPS, timestamped photo, asset ID, supervisor approval workflow"),
            ("Digitized verifiable taxpayer register", "PostgreSQL/PostGIS with UUID keys, deduplication rules, audit trail"),
            ("Centralized GIS geodatabase", "File GDB + shapefiles + REST export; ISO 19115 metadata package"),
            ("Analyse gaps, performance & leakages", "Billed vs mapped vs potential yield modelling; ward compliance heat maps"),
            ("OSR enhancement & enforcement strategy", "18-month prioritised roadmap: automation, enforcement tiers, digital payments"),
            ("Stakeholder engagement & validation", "Two county workshops + 14 ward sensitisation sessions; consensus log"),
            ("Duration: 2 calendar months", "8-week Gantt schedule with parallel field waves (Figure 3)"),
        ],
        col_widths=[2.4, 3.8],
    )
    doc.add_page_break()

    # Part III — Methodology
    part_heading(doc, "Part III", f"Technical Approach — {FRAMEWORK}")
    section(doc, "III.1 Framework Overview")
    add_para(
        doc,
        f"MAIMA's proprietary {FRAMEWORK} (Geospatial Revenue Intelligence & Digitization for OSR) is a "
        "seven-stage closed-loop methodology integrating legal compliance, participatory governance, field "
        "enumeration, geospatial engineering, economic analysis, and enforcement planning. Each stage has "
        "defined inputs, activities, quality gates, and deliverables.",
    )
    add_figure(doc, wf, f"Figure 1 — {FRAMEWORK} end-to-end workflow with quality gates.", width=6.5)

    add_table(
        doc,
        ["Stage", "GRID-OSR Phase", "Key Activities", "Quality Gate / Output"],
        [
            ("G", "Governance & Inception", "Contract signing, stakeholder mapping, inception report, Technical Committee briefing", "Inception Report — Week 1 (20% payment)"),
            ("R", "Revenue Legal Review", "Finance Act analysis, existing register audit, revenue stream taxonomy", "Approved revenue catalogue & data dictionary"),
            ("I", "Instrument (Survey123)", "Form design, pilot in 2 wards, enumerator training (36 staff)", "Data Collection Tools — Week 2"),
            ("D", "Deploy Field Teams", "4-wave ward rollout: Nyando → Muhoroni/Seme → West/East → Central fringe", "Draft Inventory & Register — Weeks 3–5 (30% payment)"),
            ("O", "Organise Geodatabase", "Topology QA, register deduplication, spatial joins, metadata", "Validated geodatabase — Week 6"),
            ("S", "Strategise OSR Plan", "Gap/leakage analysis, revenue potential, enforcement tiers, automation roadmap", "Draft Report — Week 7 (30% payment)"),
            ("R", "Report & Handover", "Validation workshops, final report, CEC presentation, ICT training", "Final deliverables — Week 8 (20% payment)"),
        ],
        col_widths=[0.4, 1.2, 2.6, 1.9],
    )

    section(doc, "III.2 Guiding Principles")
    add_bullets(
        doc,
        [
            "Legal fidelity — every mapped feature tagged to Finance Act section and applicable fee schedule",
            "Evidence-based enumeration — mandatory GPS + photo; no desk-only or proxy records accepted",
            "Participatory ownership — ward administrators co-sign cluster completion certificates",
            "Dual-capture QA — 12% independent back-check with ≤3% variance tolerance",
            "Gender & social inclusion — market stall mapping captures women/youth trader segments",
            "Interoperability — export schemas compatible with county ERP/billing system integration",
            "Data protection — compliance with Kenya Data Protection Act 2019; consent where applicable",
        ],
    )

    section(doc, "III.3 Revenue Streams Catalogue")
    add_table(
        doc,
        ["Finance Act Category", "Asset types captured", "Key attributes"],
        [
            ("Commercial", "SBP, markets (stalls/shops/slabs), bus parks, slaughterhouses, ODA, casinos, pool tables, social halls", "Permit no., trader ID, stall count, fee class, compliance status"),
            ("Public utilities", "Public toilets, weigh bridges, fish landing sites, solar storage/cooling plants", "Operator, capacity, tariff band, licence history"),
            ("Recreation & leisure", "Stadiums, swimming pools, leisure boats, ferries, water bus", "Facility type, capacity, event licence records"),
            ("Property & housing", "Rateable residential and commercial units", "Plot ref., units, occupancy, last assessment year"),
        ],
        col_widths=[1.2, 2.6, 2.4],
    )

    section(doc, "III.4 Field Enumeration Process")
    add_para(doc, "Figure 5 illustrates the three-layer field QA workflow from enumerator capture through supervisor approval to independent back-check.")
    add_figure(doc, swim, "Figure 5 — Field enumeration swimlane: enumerator, ward supervisor, and QA team.", width=6.5)

    section(doc, "III.5 Stakeholder Engagement")
    add_para(doc, "MAIMA adopts a continuous participatory cycle ensuring county ownership and validation consensus (Figure 6).")
    add_figure(doc, stake, "Figure 6 — Stakeholder engagement and validation cycle.", width=4.5)
    add_bullets(
        doc,
        [
            "Inception stakeholder mapping session with Revenue Directorate and ward administrators",
            "14 ward-level sensitisation briefings before field deployment in each ward",
            "Two county validation workshops (draft and final) with documented consensus log",
            "Market committee and trader association engagement for commercial stream verification",
            "Executive briefing to CEC Finance and Economic Planning with actionable recommendations",
        ],
    )
    doc.add_page_break()

    # Part IV — Technology
    part_heading(doc, "Part IV", f"Technology Architecture — {SOLUTION}")
    section(doc, "IV.1 Platform Overview")
    add_para(
        doc,
        f"The {SOLUTION} ({SOLUTION_FULL}) is MAIMA's integrated technology stack for county OSR assignments. "
        "It comprises field capture, core geodatabase, analytics engine, and executive reporting layers.",
    )
    add_figure(doc, arch, f"Figure 4 — {SOLUTION} layered technology architecture.", width=6.2)

    section(doc, "IV.2 Component Specifications")
    add_table(
        doc,
        ["Component", "Technology", "Function"],
        [
            ("K-REVMAP Field", "ESRI Survey123", "Offline-first mobile capture with GPS, photo, supervisor workflow"),
            ("K-REVMAP Core", "PostgreSQL/PostGIS + File GDB", "Taxpayer–asset relational model with version history"),
            ("K-REVMAP Insight", "ArcGIS Dashboard", "Real-time enumeration progress, QC flags, compliance heat maps"),
            ("K-REVMAP Analytics", "Python/R scripts", "Gap/leakage modelling, revenue potential, ward benchmarking"),
            ("K-REVMAP Export", "SHP, CSV, REST API", "County billing system integration and open data publishing"),
        ],
        col_widths=[1.3, 1.5, 3.4],
    )

    section(doc, "IV.3 Data Standards & Metadata")
    add_bullets(
        doc,
        [
            "Coordinate system: WGS 84 / UTM Zone 36S (EPSG:32736) with county boundary alignment",
            "Topology rules: no overlaps, mandatory point geometry for all revenue assets",
            "Metadata: ISO 19115 XML for geodatabase, data dictionary in Excel and PDF",
            "Naming convention: KIS_OSR_{stream}_{ward}_{YYYYMMDD} for all feature classes",
            "Backup: daily field sync to secure cloud; weekly geodatabase snapshots",
        ],
    )
    doc.add_page_break()

    # Part V — Schedule
    part_heading(doc, "Part V", "Implementation Schedule & Work Plan")
    section(doc, "V.1 Master Gantt Chart")
    add_para(
        doc,
        "Figure 3 presents the 8-week implementation schedule with parallel workstreams across inception, "
        "tooling, four field waves, analysis, validation, and handover.",
    )
    add_figure(doc, gantt, "Figure 3 — 8-week implementation Gantt chart (July–August 2026).", width=6.8)

    section(doc, "V.2 Weekly Activity Schedule")
    add_table(
        doc,
        ["Week", "GRID-OSR Phase", "Major Activities", "Milestone / Deliverable"],
        [
            ("1", "Governance", "Contract signing, inception, stakeholder mapping, Finance Act review", "Inception Report"),
            ("2", "Instrument", "Survey123 design, data dictionary, enumerator training, 2-ward pilot", "Data Collection Tools"),
            ("3", "Deploy (Wave 1)", "Nyando Belt field enumeration — Awasi, Ahero, Kabonyo, Kobura", "Daily sync dashboards"),
            ("4", "Deploy (Wave 2)", "Muhoroni & Seme corridors; agro-industrial and bus park assets", "Cluster completion certs"),
            ("5", "Deploy (Wave 3–4)", "Kisumu West/East + Central fringe wards", "Draft Inventory & Register"),
            ("6", "Organise", "Geodatabase build, topology QA, register deduplication", "Clean OSR Geodatabase"),
            ("7", "Strategise", "Gap analysis, draft report, Validation Workshop #1, sensitisation", "Draft OSR Mapping Report"),
            ("8", "Report", "Final report, Validation Workshop #2, CEC presentation, ICT training", "Final deliverables & handover"),
        ],
        col_widths=[0.6, 1.0, 2.8, 1.8],
    )

    section(doc, "V.3 Reporting & Governance")
    add_bullets(
        doc,
        [
            "Fortnightly written progress reports to Chief Officer, Finance and Economic Planning",
            "Weekly 30-minute Technical Committee stand-up (Director of Revenue as chair)",
            "Shared K-REVMAP Insight dashboard for live field progress and QC visibility",
            "Risk register maintained and updated fortnightly with mitigation actions",
            "Change log for scope adjustments requiring county written approval",
        ],
    )

    section(doc, "V.4 County Obligations Acknowledged")
    add_bullets(
        doc,
        [
            "Stakeholder facilitation and public participation sessions as needed",
            "Security for field teams in sensitive areas",
            "Provision of existing registers, Finance Act, billing records, and ward boundary data",
            "Timely review and approval of deliverables per payment milestone schedule",
            "Venue and participant mobilisation for validation workshops",
        ],
    )
    doc.add_page_break()

    # Part VI — QA & Risk
    part_heading(doc, "Part VI", "Quality Assurance & Risk Management")
    section(doc, "VI.1 Quality Assurance Framework")
    add_table(
        doc,
        ["QA Layer", "Method", "Threshold / Standard"],
        [
            ("Field — Supervisor", "100% review of first 50 assets per ward; random 10% thereafter", "Zero unresolved rejections before sync"),
            ("Field — Back-check", "Independent team re-visits 12% random sample per ward", "≤3% attribute variance tolerance"),
            ("Technical — GIS", "Automated topology, attribute, and geometry validation rules", "Zero critical errors at geodatabase gate"),
            ("Legal — Revenue", "Finance Act crosswalk review with Revenue Director", "100% stream categories signed off"),
            ("Stakeholder", "Validation workshop consensus log with action tracker", "All material corrections incorporated in final"),
            ("Document", "Peer review of all reports before submission", "Team Leader sign-off on every deliverable"),
        ],
        col_widths=[1.2, 2.8, 2.2],
    )

    section(doc, "VI.2 Risk Register")
    add_table(
        doc,
        ["Risk", "Likelihood", "Impact", "Mitigation"],
        [
            ("Weather disrupts field work", "Medium", "Medium", "Flexible wave scheduling; enumerator reserve pool of 6"),
            ("Incomplete existing registers", "High", "Medium", "Early desk audit Week 1; ward admin interviews to fill gaps"),
            ("Security in remote wards", "Low", "High", "County-provided security; field teams work in pairs; daily check-ins"),
            ("Enumerator attrition", "Medium", "Medium", "Over-recruit by 10%; cross-trained supervisors as backup enumerators"),
            ("Data sync failures", "Low", "Medium", "Offline-first Survey123; daily manual export backup protocol"),
            ("Stakeholder scheduling delays", "Medium", "Low", "Early calendar blocking at inception; virtual fallback sessions"),
            ("Scope creep beyond 14 wards", "Low", "High", "Written change control; city boundary GIS layer locked at inception"),
        ],
        col_widths=[1.6, 0.7, 0.7, 3.2],
    )
    doc.add_page_break()

    # Part VII — Team
    part_heading(doc, "Part VII", "Team Composition & Consultant Profiles")
    add_para(
        doc,
        f"MAIMA proposes a team exceeding TOR minimums: Lead Consultant, GIS Expert, 2 Revenue Analysts, "
        "Sociologist, Database Expert, QA Manager, Field Coordinator, ICT Analyst, and 36 enumerators "
        "with 14 ward supervisors. All core consultants are based in Nyanza/Western Kenya. "
        "Names below are proposed team members — final CVs attached in submission annex.",
    )
    add_table(
        doc,
        ["Consultant", "Role", "Person-days"],
        [(m["name"], m["role"], str(m["loe"])) for m in TEAM]
        + [
            ("Field enumerators (36)", "Survey123 data collection across 14 wards", "36 × 14 days"),
            ("Ward supervisors (14)", "Daily field supervision and QA", "14 × 14 days"),
        ],
        col_widths=[1.8, 2.8, 1.0],
    )

    section(doc, "VII.1 Consultant Profiles")
    for m in TEAM:
        p = doc.add_heading(m["name"], level=3)
        for r in p.runs:
            r.font.color.rgb = GREEN_RGB
            r.font.size = Pt(11)
        add_para(doc, f"Proposed role: {m['role']}", bold=True)
        add_para(doc, f"Qualifications: {m['qual']}")
        add_para(doc, f"Relevant experience: {m['exp']}")
        add_para(doc, f"Level of effort: {m['loe']} person-days.", space_after=10)
    doc.add_page_break()

    # Part VIII — Deliverables
    part_heading(doc, "Part VIII", "Deliverables & Payment Milestones")
    add_table(doc, ["No.", "Deliverable", "Payment (%)"], PAYMENT_ROWS + [("", "TOTAL (inclusive of taxes, facilitation & all costs)", "100%")], col_widths=[0.4, 4.2, 1.0])

    section(doc, "VIII.1 Detailed Deliverables Checklist")
    add_table(
        doc,
        ["#", "Deliverable", "Format", "Timing"],
        [
            ("1", "Project Inception Report", "Word + PDF (6 hard copies)", "Week 1"),
            ("2", "Survey123 forms & data dictionary", "Survey123 package + Excel + PDF", "Week 2"),
            ("3", "Draft revenue inventory & taxpayer register", "GDB + SHP + Excel", "Week 5"),
            ("4", "Draft OSR Mapping & Strategy Report", "Word + PDF (6 hard copies)", "Week 7"),
            ("5", "Validation Workshop Report #1", "Word + PDF + attendance", "Week 7"),
            ("6", "Final OSR Mapping & Strategy Report", "Word + PDF (6 hard copies)", "Week 8"),
            ("7", "Final geodatabase & metadata package", "GDB + SHP + ISO 19115 XML", "Week 8"),
            ("8", "Executive presentation to CEC", "PowerPoint + briefing note", "Week 8"),
        ],
        col_widths=[0.3, 2.4, 1.8, 0.8],
    )

    section(doc, "VIII.2 Cost Inclusions")
    add_bullets(
        doc,
        [
            "All professional fees, VAT, and applicable taxes",
            "36 enumerators and 14 ward supervisors — recruitment, training, and field allowances",
            "ESRI Survey123 enterprise licences for assignment duration",
            "Transport, accommodation, and security for field teams",
            "Two stakeholder validation/sensitisation workshops",
            "County ICT staff training on geodatabase maintenance (1-day session)",
        ],
    )
    doc.add_page_break()

    # Part IX — References
    part_heading(doc, "Part IX", "References & Track Record")
    add_para(doc, "MAIMA has completed the following similar assignments within the last five years:")
    add_table(doc, ["Project", "Client", "Year", "Scope & Outcome"], REFERENCES, col_widths=[1.6, 1.4, 0.5, 2.7])
    add_para(doc, "Reference letters and contact details for referees are provided in the submission annex.", size=10)

    # Part X — Sustainability
    part_heading(doc, "Part X", "Sustainability, Handover & Legal Compliance")
    section(doc, "X.1 Sustainability & County Ownership")
    add_bullets(
        doc,
        [
            "Geodatabase designed for county ICT staff to update without external consultant dependency",
            "One-day hands-on training for county GIS/revenue staff on Survey123 and geodatabase maintenance",
            "Standard Operating Procedures (SOP) manual for ongoing asset registration and register updates",
            "Recommended annual refresh cycle and enumerator re-training plan",
            "Open export formats (SHP, CSV) ensuring no vendor lock-in",
        ],
    )

    section(doc, "X.2 Legal & Policy Framework")
    add_bullets(
        doc,
        [
            "Constitution of Kenya, 2010",
            "County Governments Act, 2012",
            "Public Finance Management Act, 2012",
            "Urban Areas and Cities Act, 2011 (revised 2019)",
            "Kenya Data Protection Act, 2019",
            "Kisumu County Finance Act (latest); Fiscal Strategy Paper; CIDP",
        ],
    )

    section(doc, "X.3 Professional Fees")
    add_para(
        doc,
        "Detailed pricing is provided in the separate Financial Proposal. Fees are structured per TOR "
        "milestone schedule and include all costs listed in Section VIII.2.",
    )

    # Declaration
    part_heading(doc, "Declaration", "")
    add_para(
        doc,
        f"We, {COMPANY}, declare that this proposal is prepared in full compliance with the Terms of Reference, "
        "that all information provided is accurate and complete, that we have no conflict of interest, "
        "and that we accept the county's general contract conditions.",
        space_after=16,
    )
    add_para(doc, "Signature: ________________________________")
    add_para(doc, f"Name: {LEAD}")
    add_para(doc, "Title: Team Leader / Lead Consultant")
    add_para(doc, f"Date: {SUBMISSION_DATE}")
    add_para(doc, COMPANY, bold=True, color=GREEN_RGB)
    add_para(doc, ADDRESS)
    add_para(doc, EMAIL)

    doc.save(OUT)
    print(f"Wrote {OUT}")
    print(f"Diagrams: {len(diagrams)} files in {DIAG_DIR}")


if __name__ == "__main__":
    build()
