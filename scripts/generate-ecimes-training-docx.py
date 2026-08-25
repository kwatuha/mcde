#!/usr/bin/env python3
"""
Generate E-CIMES User Training Script Word documents:
  - docs/E-CIMES-User-Training-Script.docx (full script)
  - docs/E-CIMES-Training-Part1-Getting-Started.docx (teleprompter)
  - docs/E-CIMES-Training-Part2-Core-Work.docx (teleprompter)
  - docs/E-CIMES-Training-Part3-Support-Wrap-Up.docx (teleprompter)

Source of truth for video tutorials. Keep in sync with Help & Support / User Manual.
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
OUT_FULL = DOCS / "E-CIMES-User-Training-Script.docx"
OUT_PARTS = [
    {
        "path": DOCS / "E-CIMES-Training-Part1-Getting-Started.docx",
        "part": 1,
        "title": "Part 1 — Getting Started",
        "subtitle": "Login, role workspaces, and ribbon navigation",
        "chapters": (1, 5),
        "length": "~10–12 minutes",
    },
    {
        "path": DOCS / "E-CIMES-Training-Part2-Core-Work.docx",
        "part": 2,
        "title": "Part 2 — Core Work",
        "subtitle": "Dashboards, AI, projects, planning, impact, finance, monitoring",
        "chapters": (6, 12),
        "length": "~12–15 minutes",
    },
    {
        "path": DOCS / "E-CIMES-Training-Part3-Support-Wrap-Up.docx",
        "part": 3,
        "title": "Part 3 — Support & Wrap-Up",
        "subtitle": "Procurement, data, reports, public portal, roles, and help",
        "chapters": (13, 17),
        "length": "~8–10 minutes",
    },
]

NAVY = RGBColor(0x0D, 0x47, 0xA1)
BLUE = RGBColor(0x15, 0x65, 0xC0)
GREEN = RGBColor(0x2E, 0x7D, 0x32)
SLATE = RGBColor(0x33, 0x33, 0x33)

# Structured chapters (number → content). Update here; regenerate Word files.
CHAPTERS: list[dict] = [
    {
        "n": 1,
        "title": "Chapter 1 — Introduction (~1 min)",
        "voiceover": (
            "Hello and welcome to the Electronic County Integrated Monitoring and Evaluation System, "
            "known as E-CIMES, for the County Government of Machakos. "
            "E-CIMES supports the full project lifecycle. "
            "First: planning — setting county priorities and programmes. "
            "Next: registration — capturing projects in the system. "
            "Next: procurement — linking tenders and awards. "
            "Next: implementation — tracking delivery on the ground. "
            "Next: monitoring — collecting field data and reviews. "
            "Next: impact evaluation — measuring results and outcomes. "
            "Next: finance — certificates, payments, and budgets. "
            "Next: reporting — dashboards and formal reports. "
            "And finally: public transparency — sharing progress with citizens — all in one platform. "
            "After you log in, E-CIMES opens the workspace that matches your role. "
            "First: Village M&E — drafting monitoring reports and village-level reviews. "
            "Next: Ward M&E — reviewing village submissions and forwarding them up the chain. "
            "Next: Sub-County M&E — consolidating ward reviews before they reach the department. "
            "Next: Department Chief Officer — reviewing projects and publishing them for public viewing. "
            "Next: Sector M&E Champions — reviewing reports from the various departments that make up a sector. "
            "Payment certificates follow their own path across three workspaces. "
            "First: Contractor — opens the contractor dashboard, selects the linked project, and initiates a "
            "payment request for work completed on site. "
            "Next: Resident Engineer — opens the engineer workspace, creates the payment certificate from that "
            "contractor request, checks quantities and supporting documents, then forwards the certificate to the "
            "Chief Engineer for review. "
            "Next: Chief Engineer — opens the chief engineer workspace and reviews the payment certificate carefully. "
            "The Chief Engineer may approve it, or return it to the Resident Engineer for correction or amendments. "
            "When the certificate is okay, the Chief Engineer forwards the approved payment certificate to the "
            "Chief Finance Officer. "
            "Finally: Chief Finance Officer — receives the approved payment certificate in the Co-Finance workspace "
            "for final finance action. "
            "In this training you will learn how to log in, recognise your role workspace, navigate "
            "the system, work with projects and finance, measure programme impact, record monitoring data, "
            "approve certificates, generate reports, and get help when you need it."
        ),
        "on_screen": (
            "Lifecycle stills (assets/lifecycle/), then workspace screenshot montage timed to audio: "
            "M&E chain Village → Ward → Sub-County → Department Chief → Sector Champions; then payment path "
            "Contractor → Resident Engineer → Chief Engineer → Chief Finance Officer. "
            "See assets/workspace-timing.csv."
        ),
        "tips": [
            "Features visible to each user depend on role, permissions, organisation scope, and UI profile. "
            "If you cannot see a menu item described here, contact your ICT administrator."
        ],
    },
    {
        "n": 2,
        "title": "Chapter 2 — Logging In (~2 min)",
        "voiceover": (
            "Open a supported browser — Chrome, Edge, or Firefox — and go to the E-CIMES staff URL "
            "provided by ICT. Enter your username or email and password. If your account has OTP enabled, "
            "enter the verification code sent to your registered phone or email. If this is your first login "
            "or your password has expired, the system will ask you to set a new password before you continue. "
            "Click Sign In. After a successful login, E-CIMES redirects you automatically to your role landing "
            "page — you do not need to hunt for the right home screen."
        ),
        "on_screen": (
            "Login page with county branding; OTP screen if applicable; forced password change; then role landing."
        ),
        "bullets": [
            "Do not share passwords or OTP codes.",
            "If login fails, confirm your account is active and approved.",
            "Missing menus after login usually means role, scope, or UI profile — not a system error.",
            "Contractors must be linked in Contractor Management; otherwise the portal shows “Account not linked yet”.",
        ],
    },
    {
        "n": 3,
        "title": "Chapter 3 — After Login: Your Role Workspace (~3 min)",
        "voiceover": (
            "E-CIMES no longer sends every user to the same Personal Dashboard. Your landing page and sidebar "
            "depend on your role. Portal roles — Village Administrator, Ward Administrator, Sub-County "
            "Administrator, Department Chief Officer, Sector M&E Champion, Resident Engineer, Chief Engineer, "
            "Chief Finance Officer, and Contractor — open a focused Workspace with a short sidebar of the tools "
            "they need. County leadership roles such as Executive Supervisor and Executive Viewer open Summary "
            "Statistics for briefing dashboards. Other staff — operators, reviewers, and ICT — open the Personal "
            "Dashboard and use the full ribbon menu. Look at the sidebar group title after login: Village M&E, "
            "Ward M&E, Sub-County M&E, Department Chief M&E, Sector M&E Champions, Resident Engineer, Chief "
            "Engineer, Co-Finance — which means Chief Finance Officer — or Contractor. That title tells you which workspace you are in. For "
            "self-service guidance, open the three-dot menu (top right) and select Help & Support."
        ),
        "on_screen": (
            "Split-screen or sequence of landings: Village M&E Workspace, Ward M&E Workspace, Resident Engineer "
            "Workspace, Summary Statistics, Personal Dashboard with full ribbon."
        ),
        "table": {
            "headers": ["Role / audience", "Landing after login", "Navigation style"],
            "rows": [
                ["Village Administrator", "Village M&E Workspace", "Short sidebar: Village M&E"],
                ["Ward Administrator", "Ward M&E Workspace", "Short sidebar: Ward M&E"],
                ["Sub-County Administrator", "Sub-County M&E Workspace", "Short sidebar: Sub-County M&E"],
                ["Department Chief Officer", "Department Chief M&E Workspace", "Short sidebar: Department Chief M&E"],
                ["Sector M&E Champion", "Sector M&E Champions Workspace", "Short sidebar: Sector M&E Champions"],
                ["Resident Engineer", "Resident Engineer Workspace", "Short sidebar: Workspace tools"],
                ["Chief Engineer", "Chief Engineer Workspace", "Short sidebar: Workspace tools"],
                ["Chief Finance Officer (Co-Finance)", "Co-Finance Workspace", "Short sidebar: Co-Finance"],
                ["Contractor", "Contractor Dashboard", "Short sidebar: Contractor"],
                ["Governor / Executive Supervisor or Viewer", "Summary Statistics", "Dashboard ribbon (briefing)"],
                ["Operators, reviewers, ICT / Admin", "Personal Dashboard", "Full ribbon menu"],
            ],
        },
        "tips": [
            "Check Notifications and Approvals on Personal Dashboard when you have access to it — pending users, "
            "projects, reviews, PMC reports, and workflow tasks appear there."
        ],
    },
    {
        "n": 4,
        "title": "Chapter 4 — Role Workspaces in Detail (~5 min)",
        "voiceover": (
            "This chapter walks through the dedicated workspaces. Trainers may pause and demo only the roles in the room."
        ),
        "sections": [
            {
                "heading": "4.1 M&E monitoring chain — Village → Ward → Sub-County → Chief",
                "voiceover": (
                    "Village Administrators land on Village M&E Workspace. Use summary cards for Draft, With ward, "
                    "Approved, and Returned. Open My drafts to complete checklist reports and progress status, then "
                    "submit to the ward. Attach photos and documents from the workspace quick actions or My projects. "
                    "Village, Ward, Sub-County, and Chief Finance Officer workspaces also show Projects by department so you can "
                    "see which departments own which projects in your scope. Ward Administrators land on Ward M&E "
                    "Workspace. Use the Ward review queue to revise, return, or forward village reports; Village drafts "
                    "and All reports tabs show the full pipeline. Sub-County Administrators review ward-forwarded "
                    "reports on Sub-County M&E Workspace — return with comments or forward to the Department Chief. "
                    "Department Chief Officers use Department Chief M&E Workspace for final approval; approving publishes "
                    "the project toward the citizen dashboard. Sector M&E Champions use a read-only Sector M&E Champions "
                    "Workspace to oversee All sector reports and Pending chief approval across mapped departments. "
                    "The chain is: Village draft → Ward review → Sub-county review → Chief approve → public."
                ),
                "on_screen": (
                    "Village workspace My drafts → submit; Ward review queue → forward; Sub-county review → forward; "
                    "Chief approval queue → approve; optional Projects by department card."
                ),
                "bullets": [
                    "Village sidebar — Workspace, Monitoring reports, Monitoring visits, My projects, Project documents",
                    "Ward sidebar — Workspace, Ward review queue, Village drafts, Ward projects, Project documents",
                    "Sub-County sidebar — Workspace, Sub-county review queue, All reports, Sub-county projects, Project documents",
                    "Chief sidebar — Workspace, Chief approval queue, All reports, Department projects, Project documents",
                    "Sector M&E — Workspace, All sector reports, Pending chief approval, Sector projects, Project documents",
                ],
            },
            {
                "heading": "4.2 Engineering and Chief Finance Officer workspaces",
                "voiceover": (
                    "Resident Engineers open Resident Engineer Workspace for projects in their scope, progress photos, "
                    "payment requests, and first-step payment certificate approval. Chief Engineers open Chief Engineer "
                    "Workspace for second-step certificate approval after the Resident Engineer. The Chief Finance Officer "
                    "opens the Co-Finance Workspace — labelled Co-Finance in the sidebar, meaning Chief Finance Officer — "
                    "for final certificate sign-off, payment requests, project registry, and county finance tools. "
                    "Certificate order is always: Resident Engineer, then Chief Engineer, then Chief Finance Officer. "
                    "Anyone can later confirm authenticity under Finance → Verify Certificate using the QR code or "
                    "certificate number — verification also works without logging in."
                ),
                "on_screen": (
                    "Resident Engineer Certificates (step 1) → Chief Engineer Payment Certificates (step 2) → "
                    "Co-Finance (Chief Finance Officer) Payment certificates (step 3) → Verify Certificate with QR."
                ),
                "bullets": [
                    "Resident Engineer — Workspace, Project Registry, Progress Photos, Payment Requests, Certificates",
                    "Chief Engineer — Workspace, Payment Certificates, Project Registry, Payment Requests, Progress Photos",
                    "Chief Finance Officer (Co-Finance sidebar) — Workspace, Payment Certificates, Payment Requests, Project Registry, County Finance",
                ],
            },
            {
                "heading": "4.3 Contractor portal",
                "voiceover": (
                    "Contractors open Contractor Dashboard. From the Contractor sidebar, request payments, upload "
                    "progress photos, submit project files, and track payment status for assigned projects only. If "
                    "the account is not linked to a contractor record, contact ICT or Procurement before continuing."
                ),
                "on_screen": (
                    "Contractor Dashboard quick actions: Request payment, Upload progress photos, Submit project files."
                ),
            },
            {
                "heading": "4.4 County leadership briefing",
                "voiceover": (
                    "Governors and other executive supervisor or viewer roles open Summary Statistics. Use the "
                    "Dashboard ribbon for Project By Status, Finance Dashboard, GIS Map, and related analytical views. "
                    "Apply filters before interpreting numbers. On the mobile field app, executive roles use Briefing "
                    "mode rather than field checklist forms."
                ),
                "on_screen": (
                    "Summary Statistics landing; open Dashboard → Project By Status; optional mobile Briefing mode."
                ),
            },
        ],
    },
    {
        "n": 5,
        "title": "Chapter 5 — Full Ribbon Navigation (~2 min)",
        "voiceover": (
            "Staff who land on Personal Dashboard use the top ribbon for the main modules. Depending on your role "
            "and UI profile, you may also see Admin and other options. Portal workspace users usually see a short "
            "sidebar instead of the full ribbon — that is expected. On the home page, review Notifications and "
            "Approvals regularly."
        ),
        "on_screen": "Home page; open ribbon and hover each tab; open Notifications; open Help & Support.",
        "bullets": [
            "Dashboard — executive and analytical views, Personal Dashboard, Mobile app (Android)",
            "Finance — payments, certificates, verification",
            "Projects — project registry, implementation, evaluation, engineer workspaces (if permitted)",
            "Planning — CIDP, ADP, programmes, indicators, programme progress, budget alignment",
            "Data — import tools and beneficiary registry",
            "Procurement — procurement stages and contractors",
            "Monitoring — visits, PMC, village workflow, checklists, field data",
            "Reports — built-in county reports and report library",
            "Public — citizen-facing content and approvals",
            "Admin — users, metadata, audit, UI profiles (for authorised staff)",
        ],
    },
    {
        "n": 6,
        "title": "Chapter 6 — Dashboards (~2 min)",
        "voiceover": (
            "Select Dashboard from the ribbon — or open Summary Statistics if that is your landing — to use "
            "analytical views. Each dashboard answers a different question. Apply filters before interpreting "
            "numbers. Dashboard totals reflect your access scope and data completeness. Jobs & Impact shows "
            "jobs and beneficiary volumes; for CIDP and ADP outcome scorecards, also use Planning → CIDP or "
            "ADP Programme Progress, covered in the impact chapter."
        ),
        "on_screen": "Dashboard submenu; open Project By Status; demonstrate filters.",
        "table": {
            "headers": ["Dashboard", "Purpose"],
            "rows": [
                ["Personal Dashboard", "Your scoped summary and workflow inbox"],
                ["Summary Statistics", "County-wide summary cards and trends (leadership landing)"],
                ["Project By Status", "Projects grouped by implementation status"],
                ["Project By Sector", "Distribution by sector/programme"],
                ["Finance Dashboard", "Budget, disbursement, absorption KPIs"],
                ["Operations Dashboard", "Operational delivery and attention items"],
                ["Jobs & Impact", "Employment and beneficiary volume (complement with Programme Progress for outcomes)"],
                ["Regional Breakdown", "Subcounty and ward distribution"],
                ["Departmental Reports", "Department-level summaries"],
                ["GIS Dashboard / Project GIS Map", "Location-based project views"],
            ],
        },
        "tips": ["If figures look wrong, compare dashboard filters with Projects Registry filters."],
    },
    {
        "n": 7,
        "title": "Chapter 7 — AI Assistant & Help (~2 min)",
        "voiceover": (
            "E-CIMES includes an AI Assistant on every authenticated page. Click the sparkle button at the "
            "bottom-right. You can ask how-to questions, request live summaries of projects you can access, "
            "or generate professional Word or PDF reports from the screen you are on. For formal documents, "
            "click Generate Professional Report, choose the report type and format, and download. AI output "
            "is advisory — review before official use. Open the relevant workspace, dashboard, or project first "
            "so the AI uses the correct context. Help & Support contains the full user manual, including how "
            "to measure project and programme impact."
        ),
        "on_screen": (
            "Click sparkle button; ask a navigation question; show Generate Professional Report dialog."
        ),
    },
    {
        "n": 8,
        "title": "Chapter 8 — Projects Module (~2 min)",
        "voiceover": (
            "Open Projects from the ribbon, or My projects / Project Registry from your workspace sidebar. Use "
            "Projects Registry to search and filter by name, department, financial year, sector, status, "
            "subcounty, or ward. Always search before creating a new project to avoid duplicates. Open a project "
            "to review tabs: overview, milestones, status updates, documents, photos, teams, partners, funding, "
            "and certificates. Attach evidence that supports reported progress. Some tabs may be hidden by your "
            "UI profile — that is intentional. Use Projects → Project Evaluation for structured evaluation lines "
            "with baseline, target, achieved, reporting period, and result level — we cover that next under impact."
        ),
        "on_screen": (
            "Projects Registry; open project details; scroll tabs; show document upload or status form."
        ),
        "bullets": [
            "Implementation Plans — cross-project planning views",
            "Project Status / Updates — progress recording",
            "Project Evaluation — baseline → achieved, result level, reporting period",
            "Schedule & Milestones, Teams, Partners, Activity links",
        ],
    },
    {
        "n": 9,
        "title": "Chapter 9 — Planning Module (~2 min)",
        "voiceover": (
            "Select Planning to manage county planning structures: CIDP periods and pillars, ADP periods and "
            "ADP Implementation, RRI Programmes, Budget Traceability, ADP–Budget linkage, and indicator, "
            "activity, and risk catalogues. Open Indicators & KPIs to set each KPI’s result level — Output, "
            "Outcome, or Impact. Open CIDP Programme Progress and ADP Programme Progress to see delivery "
            "roll-ups and impact scorecards for linked programmes. Planning catalogues feed monitoring and "
            "reports. Portal-only roles may not see Planning — use your workspace tools instead."
        ),
        "on_screen": (
            "Planning menu; Indicators & KPIs result level; CIDP Programme Progress; ADP Implementation gap summary."
        ),
    },
    {
        "n": 10,
        "title": "Chapter 10 — Measuring Project & Programme Impact (~4 min)",
        "voiceover": (
            "Progress tells you how far implementation has gone — for example percentage complete. Impact "
            "asks whether intended benefits and outcomes are being realized for people. E-CIMES supports both. "
            "First, open Planning → Indicators & KPIs and tag each KPI as Output for delivery, Outcome for "
            "change for people, or Impact for longer-term effects. Second, link activities and indicators to "
            "projects with baseline and target values. Third, open Projects → Project Evaluation. Select a "
            "project and add evaluation lines: evaluation date, activity or indicator, baseline, target, "
            "achieved value, optional performance score, reporting period — Baseline, Midline, or Endline — "
            "and result level. Capture Midline and Endline against the same indicators used at Baseline so "
            "change is comparable. Fourth, open Data → Beneficiary Registry and set Outcome status as people "
            "move from enrolled to receiving benefit to benefit realized. Fifth, on Village Field Monitoring — "
            "web or CIMES Mobile — answer Community benefit realized and Access improved, and add short impact "
            "notes from community comments. Work stages on the checklist come from the project’s milestones — "
            "sync the mobile app after template updates. Finally, open Planning → CIDP Programme Progress or "
            "ADP Programme Progress. Review delivery columns — projects, linkage, average progress, stalled, "
            "budgets — and impact columns — Outcome achievement percent, Outcome lines, Benefit realized, and "
            "Community benefit. Export Excel or PDF for leadership packs. Remember: a finished project can still "
            "have weak community benefit — do not treat percentage complete alone as impact."
        ),
        "on_screen": (
            "Indicators result level dialog → Project Evaluation add line (period + result level) → Beneficiary "
            "outcome status → Village community Yes/No fields → CIDP Programme Progress impact columns → optional "
            "Excel export."
        ),
        "tips": [
            "If impact columns show dashes, link projects to CIDP/ADP programmes and add evaluations, beneficiary "
            "outcomes, or village community answers.",
            "Hard-refresh the browser after releases; on mobile, Sync Checklists after ICT updates the village template.",
            "Jobs & Impact Dashboard complements this chapter for volumes; Programme Progress is the outcome scorecard.",
        ],
        "bullets": [
            "Progress = delivery (% complete, milestones, certificates)",
            "Impact = outcomes for people (result levels, Baseline→Endline, benefit realized, community Yes/No)",
            "Menu paths: Planning → Indicators; Projects → Project Evaluation; Planning → CIDP/ADP Programme Progress",
        ],
    },
    {
        "n": 11,
        "title": "Chapter 11 — Financial Tracking & Certificates (~3 min)",
        "voiceover": (
            "Finance staff and Chief Finance Officer workspace users review Finance Dashboard, Payment List, Payment "
            "Certificates, Funding Sources Report, and Verify Certificate. Engineers and the Chief Finance Officer approve "
            "payment certificates inside their workspaces as covered earlier. Authorised users create interim "
            "or final certificates on the project Certificates tab. Generated PDFs include a QR code labelled "
            "Scan to verify this certificate. To verify: open Finance → Verify Certificate, scan the QR code "
            "or type the certificate number. Verification works without logging in."
        ),
        "on_screen": (
            "Finance submenu or Co-Finance (Chief Finance Officer) workspace; Certificates PDF with QR; "
            "Verify Certificate page with valid result."
        ),
        "table": {
            "headers": ["Screen", "Purpose"],
            "rows": [
                ["Finance Dashboard", "High-level finance KPIs"],
                ["Payment List", "Payment records and filters"],
                ["Payment Certificates", "County-wide certificate list"],
                ["Funding Sources Report", "Funding analysis"],
                ["Verify Certificate", "Confirm certificate authenticity (QR or manual)"],
                ["Resident / Chief Engineer / Chief Finance Officer (Co-Finance) workspaces", "Role-step certificate approval chain"],
            ],
        },
    },
    {
        "n": 12,
        "title": "Chapter 12 — Monitoring & Field Data (~3 min)",
        "voiceover": (
            "M&E portal users should prefer their role workspace for day-to-day review and approval. From the "
            "Monitoring ribbon, staff can also open Monitoring Visits, PMC Ward Reports, Village monitoring "
            "workflow, Ward Accountability, Checklists & visits, Evaluation Workbench, and Stakeholder feedback. "
            "Village Field Monitoring loads work stages from the project record and includes community impact "
            "questions used in programme scorecards. Field staff download CIMES Mobile from Dashboard → Mobile "
            "app (Android) or the mobile quick action on Village workspace, sign in, sync checklists, select a "
            "project, complete the form, and submit — offline if needed. Executive roles on mobile use Briefing mode."
        ),
        "on_screen": (
            "Monitoring menu or Ward workspace queue; Checklists & visits with community impact fields; Mobile app "
            "download page; phone with collector app."
        ),
    },
    {
        "n": 13,
        "title": "Chapter 13 — Procurement & Data (~2 min)",
        "voiceover": (
            "Procurement tracks cases from planning through award: Project Procurement, Budget Procurement "
            "Intake, Procured Projects, and Contractor Registry. Link contractor accounts so the Contractor "
            "Dashboard works. On Budget Intake, Setup scope before Quote. Data supports Import Data with "
            "template preview and validation, Data Import Logs, and Beneficiary Registry — including outcome "
            "status for impact tracking. Always fix validation errors before confirming an import."
        ),
        "on_screen": (
            "Project Procurement stages; Import Data template download and preview; Beneficiary outcome status; "
            "Contractor Registry link note."
        ),
    },
    {
        "n": 14,
        "title": "Chapter 14 — Reports (~2 min)",
        "voiceover": (
            "Open Reports and start at Reports hub — a searchable index of standard county reports. Set "
            "filters, review totals, and export. Use Report Library for approved archives and Scheduled "
            "Reports for recurring email delivery. Built-in reports use county templates; AI Professional "
            "Reports from the sparkle button produce Word/PDF drafts from the current screen or workspace. "
            "CIDP and ADP Programme Progress also export Excel and PDF with impact columns."
        ),
        "on_screen": "Reports hub search; open one report with filters and export.",
        "bullets": [
            "County Operations, APR Reports, Status Report, Absorption Report",
            "Pending Bills, Project Finance Overview, Yearly Trends",
            "Budget Justification, Reporting Template (Word), PMC Ward Reports",
            "CIDP / ADP Programme Progress exports (delivery + impact)",
        ],
    },
    {
        "n": 15,
        "title": "Chapter 15 — Public Portal (~1 min)",
        "voiceover": (
            "Staff with permission use Public Approval to release projects for the citizen dashboard, "
            "moderate feedback, and manage announcements. Projects approved through the Department Chief "
            "M&E Workspace feed into public visibility. Only approved content appears on public pages."
        ),
        "on_screen": "Public menu → Public Approval (brief); optional citizen dashboard glimpse.",
    },
    {
        "n": 16,
        "title": "Chapter 16 — Roles, Scope & Getting Help (~2 min)",
        "voiceover": (
            "What you can see and do depends on your role, organisation scope — ward, sub-county, department, "
            "sector — and UI profile. When something seems wrong, confirm you are on the expected workspace "
            "after login, clear filters, open Help & Support or ask the AI Assistant, then contact ICT with "
            "your role name, the module or workspace name, the error message, and a screenshot (no passwords)."
        ),
        "table": {
            "headers": ["Role", "Typical focus"],
            "rows": [
                ["County leadership (Executive Supervisor / Viewer)", "Summary Statistics, Dashboard ribbon, AI summaries, mobile Briefing"],
                ["CECM / County Viewer / Chief Officers (view)", "Scoped dashboards and department oversight"],
                ["Village / Ward / Sub-County / Chief M&E", "Dedicated M&E workspace and approval chain"],
                ["Sector M&E Champion", "Read-only sector portfolio and pending chief items"],
                ["Resident / Chief Engineer", "Engineer workspace: photos, payments, certificate steps 1–2"],
                ["Chief Finance Officer", "Co-Finance workspace (Chief Finance Officer): final certificates and finance tools"],
                ["Contractor", "Contractor Dashboard: payments, photos, project files"],
                ["Department / County operators & reviewers", "Projects, progress, evidence, reviews via full ribbon"],
                ["Planning / M&E staff", "Indicators, programme progress, monitoring, evaluation"],
                ["Finance / Procurement staff", "Module ribbon tools matching their function"],
                ["Field collectors", "Mobile Field Collector, offline visits, community impact questions"],
                ["ICT / Admin", "Users, UI profiles, metadata, audit trail, mobile releases"],
            ],
        },
    },
    {
        "n": 17,
        "title": "Chapter 17 — Conclusion (~1 min)",
        "voiceover": (
            "You have now seen how to log in and land on the correct role workspace; navigate M&E, engineer, "
            "Chief Finance Officer, and contractor portals; use dashboards and the full ribbon when available; manage "
            "projects and planning; measure programme impact with indicators, evaluations, beneficiaries, and "
            "CIDP/ADP scorecards; complete finance including the three-step certificate chain and QR "
            "verification; record monitoring data and use the mobile field collector; generate built-in reports "
            "and AI professional reports; and find help through Help & Support and the AI Assistant. Continue "
            "exploring modules and workspace actions relevant to your role. Thank you for watching."
        ),
        "on_screen": "Help & Support link, ICT contact, montage of role workspaces, county logo.",
    },
]


def set_doc_defaults(doc: Document, *, body_size: int = 11) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(body_size)
    for section in doc.sections:
        section.top_margin = Inches(0.85)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            table.rows[ri + 1].cells[ci].text = val
    doc.add_paragraph()


def add_voiceover(doc: Document, text: str, *, large: bool = False) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10 if large else 6)
    label = p.add_run("Voiceover: ")
    label.bold = True
    label.font.color.rgb = BLUE
    label.font.size = Pt(14 if large else 11)
    body = p.add_run(text)
    body.font.size = Pt(14 if large else 11)
    body.font.name = "Calibri"


def add_on_screen(doc: Document, text: str, *, large: bool = False) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10 if large else 6)
    label = p.add_run("ON SCREEN: ")
    label.bold = True
    label.font.color.rgb = GREEN
    label.font.size = Pt(12 if large else 11)
    body = p.add_run(text)
    body.font.size = Pt(12 if large else 11)
    body.italic = True


def add_tip(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run("Tip: ")
    r.bold = True
    p.add_run(text)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def render_chapter(doc: Document, ch: dict, *, teleprompter: bool = False) -> None:
    doc.add_heading(ch["title"], level=1)
    if ch.get("voiceover"):
        add_voiceover(doc, ch["voiceover"], large=teleprompter)
    if ch.get("on_screen"):
        add_on_screen(doc, ch["on_screen"], large=teleprompter)
    for tip in ch.get("tips") or []:
        add_tip(doc, tip)
    if ch.get("bullets"):
        add_bullets(doc, ch["bullets"])
    if ch.get("table"):
        add_table(doc, ch["table"]["headers"], ch["table"]["rows"])
    for sec in ch.get("sections") or []:
        doc.add_heading(sec["heading"], level=2)
        if sec.get("voiceover"):
            add_voiceover(doc, sec["voiceover"], large=teleprompter)
        if sec.get("on_screen"):
            add_on_screen(doc, sec["on_screen"], large=teleprompter)
        if sec.get("bullets"):
            add_bullets(doc, sec["bullets"])


def chapters_in_range(start: int, end: int) -> list[dict]:
    return [c for c in CHAPTERS if start <= c["n"] <= end]


def add_full_title(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("E-CIMES User Training Script")
    run.bold = True
    run.font.size = Pt(26)
    run.font.color.rgb = NAVY

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run(
        "Electronic County Integrated Monitoring and Evaluation System\n"
        "County Government of Machakos"
    )
    r.font.size = Pt(14)
    r.font.color.rgb = SLATE

    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    nr = note.add_run(
        "Video tutorial voiceover script with on-screen cues\n"
        "Suggested length: 28–35 minutes as one video, or three parts (~10–15 min each)\n"
        "Part 1: Ch 1–5 · Part 2: Ch 6–12 · Part 3: Ch 13–17"
    )
    nr.italic = True
    nr.font.size = Pt(10)
    doc.add_page_break()


def add_production_notes(doc: Document) -> None:
    doc.add_heading("Production Notes (for video team)", level=1)
    add_table(
        doc,
        ["Item", "Suggestion"],
        [
            [
                "Length",
                "~30 min single video, or Part 1 (Ch 1–5) + Part 2 (Ch 6–12) + Part 3 (Ch 13–17)",
            ],
            [
                "Teleprompter files",
                "Use E-CIMES-Training-Part1/2/3-*.docx — larger Voiceover text for reading while recording",
            ],
            [
                "Demo accounts",
                "Prepare one account per major workspace: Village, Ward, Resident Engineer, Chief Finance Officer (Co-Finance), "
                "Contractor, Executive, Admin (Planning/M&E for impact chapter)",
            ],
            [
                "Must-demo",
                "Role landing after login; M&E approval chain; certificate steps 1–3; QR verify; AI sparkle; "
                "Reports hub; mobile APK; Project Evaluation + CIDP Programme Progress impact columns",
            ],
            [
                "Captions",
                "On-screen labels for paths e.g. Ward M&E → Ward review queue; Planning → CIDP Programme Progress",
            ],
            [
                "B-roll",
                "GIS map, checklist on phone, PDF with QR, workspace sidebar titles, programme progress grid",
            ],
            [
                "Audience cuts",
                "For role-specific sessions, keep Ch 1–3 + matching subsection of Ch 4 + shared Ch 7/10/11/12/17",
            ],
            [
                "Screenshots",
                "Reuse docs/manual-screenshots/; capture impact screens listed in that folder’s README",
            ],
        ],
    )
    p = doc.add_paragraph()
    r = p.add_run(
        "Recording tip: Record one chapter at a time in OBS (1080p, 30 fps). Edit in CapCut or DaVinci Resolve. "
        "Always add captions. Publish as three unlisted videos matching the Part 1–3 files."
    )
    r.italic = True
    r.font.size = Pt(10)


def build_full_document() -> Document:
    doc = Document()
    set_doc_defaults(doc)
    add_full_title(doc)
    for ch in CHAPTERS:
        render_chapter(doc, ch, teleprompter=False)
    add_production_notes(doc)
    return doc


def build_part_document(meta: dict) -> Document:
    start, end = meta["chapters"]
    chapters = chapters_in_range(start, end)
    doc = Document()
    set_doc_defaults(doc, body_size=14)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"E-CIMES Video Teleprompter\n{meta['title']}")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = NAVY

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run(
        f"{meta['subtitle']}\n"
        f"Chapters {start}–{end} · Suggested length: {meta['length']}\n"
        "Read the Voiceover aloud. Follow ON SCREEN cues while recording."
    )
    r.font.size = Pt(12)
    r.font.color.rgb = SLATE

    how = doc.add_paragraph()
    how.paragraph_format.space_before = Pt(8)
    hr = how.add_run(
        "How to use: Open this file beside OBS. Pause between chapters. Re-record a chapter if you stumble. "
        "Speak slower than feels natural. Enlarge the browser to 110–125%."
    )
    hr.italic = True
    hr.font.size = Pt(11)

    doc.add_page_break()

    for ch in chapters:
        render_chapter(doc, ch, teleprompter=True)
        # Soft page break between long chapters for easier scrolling while recording
        if ch["n"] in {4, 10, 12}:
            doc.add_page_break()

    end = doc.add_paragraph()
    end.paragraph_format.space_before = Pt(18)
    er = end.add_run(
        f"END OF {meta['title'].upper()}\n"
        "Save the recording, then edit captions in CapCut / Resolve before publishing."
    )
    er.bold = True
    er.font.color.rgb = NAVY
    return doc


def main() -> None:
    DOCS.mkdir(parents=True, exist_ok=True)

    full = build_full_document()
    full.save(OUT_FULL)
    print(f"Wrote {OUT_FULL} ({len(CHAPTERS)} chapters)")

    for meta in OUT_PARTS:
        part_doc = build_part_document(meta)
        part_doc.save(meta["path"])
        start, end = meta["chapters"]
        print(f"Wrote {meta['path'].name} (Ch {start}–{end})")


if __name__ == "__main__":
    main()
