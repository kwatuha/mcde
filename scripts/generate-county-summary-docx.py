#!/usr/bin/env python3
"""Generate formatted Word one-pager for E-CIMES county outreach."""

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

OUTPUT = Path(__file__).resolve().parents[1] / "docs" / "E-CIMES-County-One-Page-Summary.docx"

# County / Kenya-inspired professional palette
BRAND_DARK = RGBColor(0x0B, 0x3D, 0x2E)   # deep green
BRAND_ACCENT = RGBColor(0x1B, 0x7A, 0x4E)  # accent green
BRAND_GOLD = RGBColor(0xC9, 0xA2, 0x27)     # gold accent
TEXT_DARK = RGBColor(0x1A, 0x1A, 0x1A)
TEXT_MUTED = RGBColor(0x4A, 0x4A, 0x4A)
TABLE_HEADER_BG = "0B3D2E"
TABLE_ALT_BG = "F4F8F6"


def set_cell_shading(cell, fill_hex: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill_hex)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def set_paragraph_spacing(paragraph, before=0, after=6, line=1.15):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def add_run(paragraph, text, *, bold=False, italic=False, size=11, color=TEXT_DARK, font_name="Calibri"):
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.font.name = font_name
    return run


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    sizes = {1: 16, 2: 13, 3: 11.5}
    colors = {1: BRAND_DARK, 2: BRAND_ACCENT, 3: BRAND_DARK}
    add_run(p, text, bold=True, size=sizes.get(level, 11), color=colors.get(level, BRAND_DARK))
    set_paragraph_spacing(p, before=10 if level > 1 else 4, after=6)
    return p


def add_body(doc, text, *, bold_prefix=None, after=6):
    p = doc.add_paragraph()
    if bold_prefix:
        add_run(p, bold_prefix, bold=True, size=10.5)
        add_run(p, text, size=10.5, color=TEXT_MUTED)
    else:
        add_run(p, text, size=10.5, color=TEXT_MUTED)
    set_paragraph_spacing(p, after=after)
    return p


def add_bullet(doc, bold_part, rest):
    p = doc.add_paragraph(style="List Bullet")
    add_run(p, bold_part, bold=True, size=10.5, color=TEXT_DARK)
    add_run(p, rest, size=10.5, color=TEXT_MUTED)
    set_paragraph_spacing(p, after=3)


def style_table_header(row):
    for cell in row.cells:
        set_cell_shading(cell, TABLE_HEADER_BG)
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.size = Pt(9.5)
            if not paragraph.runs:
                run = paragraph.add_run(paragraph.text)
                paragraph.text = ""
                run.font.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.size = Pt(9.5)


def fill_table_cell(cell, text, *, bold=False, size=9.5):
    cell.text = ""
    p = cell.paragraphs[0]
    add_run(p, text, bold=bold, size=size, color=TEXT_DARK if bold else TEXT_MUTED)
    set_paragraph_spacing(p, after=2, before=2)


def build_document() -> Document:
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    # Title block
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(title, "E-CIMES", bold=True, size=28, color=BRAND_DARK, font_name="Calibri Light")
    set_paragraph_spacing(title, after=2)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(
        subtitle,
        "County Integrated Monitoring and Evaluation System",
        bold=True,
        size=13,
        color=BRAND_ACCENT,
    )
    set_paragraph_spacing(subtitle, after=4)

    tagline = doc.add_paragraph()
    tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(
        tagline,
        "One platform to plan, deliver, monitor, finance, and publicly account for county investments",
        italic=True,
        size=11,
        color=TEXT_MUTED,
    )
    set_paragraph_spacing(tagline, after=10)

    # Divider line via border paragraph
    divider = doc.add_paragraph()
    divider.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(divider, "―" * 42, size=8, color=BRAND_GOLD)
    set_paragraph_spacing(divider, after=10)

    add_heading(doc, "The challenge counties face", level=2)
    add_body(
        doc,
        "County governments manage hundreds of projects across departments, wards, and financial years—often "
        "tracked in spreadsheets, disconnected systems, and paper files. Leadership lacks real-time visibility. "
        "Field teams struggle to collect evidence. Finance and engineering workflows run in parallel silos. "
        "Citizens cannot easily see what is being built, where money goes, or how to give feedback.",
    )
    p = doc.add_paragraph()
    add_run(p, "E-CIMES", bold=True, size=10.5, color=BRAND_DARK)
    add_run(
        p,
        " closes that gap: a proven, end-to-end platform that connects planning, procurement, implementation, "
        "monitoring, finance, and public transparency in one secure system.",
        size=10.5,
        color=TEXT_MUTED,
    )
    set_paragraph_spacing(p, after=8)

    add_heading(doc, "What E-CIMES delivers", level=2)

    deliver_table = doc.add_table(rows=5, cols=3)
    deliver_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    deliver_table.autofit = True
    headers = [
        "For leadership",
        "For departments & M&E",
        "For citizens",
    ]
    rows_data = [
        headers,
        [
            "Executive dashboards and briefing views",
            "Role-based workspaces (Village → Ward → Sub-County → Chief)",
            "Public Investment Portal with project maps and galleries",
        ],
        [
            "Finance absorption, payment gaps, and at-risk projects",
            "Projects Registry with milestones, photos, documents, and GIS",
            "Citizen feedback and project evaluation",
        ],
        [
            "Regional breakdown by sub-county and ward",
            "Monitoring checklists with approval workflows",
            "County-proposed projects and announcements",
        ],
        [
            "AI-assisted summaries and professional reports",
            "CIDP/ADP planning, indicators, and programme scorecards",
            "Certificate verification via QR code",
        ],
    ]
    for r_idx, row_data in enumerate(rows_data):
        row = deliver_table.rows[r_idx]
        for c_idx, text in enumerate(row_data):
            fill_table_cell(row.cells[c_idx], text, bold=(r_idx == 0))
            if r_idx == 0:
                for paragraph in row.cells[c_idx].paragraphs:
                    paragraph.text = ""
                    add_run(paragraph, text, bold=True, size=9.5, color=RGBColor(0xFF, 0xFF, 0xFF))
        if r_idx == 0:
            style_table_header(row)
        elif r_idx % 2 == 0:
            for cell in row.cells:
                set_cell_shading(cell, TABLE_ALT_BG)

    doc.add_paragraph()

    add_heading(doc, "Core capabilities", level=2)
    capabilities = [
        (
            "Strategic planning & accountability",
            " — Link CIDP and ADP programmes to budget lines, projects, indicators, and risks. Track delivery and impact (outputs, outcomes, community benefit) in programme progress scorecards.",
        ),
        (
            "Full project lifecycle",
            " — From budget procurement intake through contractor assignment, implementation, evaluation, and closure—with audit trails at every step.",
        ),
        (
            "Field monitoring that works offline",
            " — CIMES Mobile (Android) lets field staff collect geotagged photos, GPS, and checklist data offline, then sync when connected. Executives get a mobile briefing dashboard on the same app.",
        ),
        (
            "Finance & payment integrity",
            " — Payment certificates follow a structured Resident Engineer → Chief Engineer → Co-Finance approval chain. Every certificate PDF includes a QR code for public verification.",
        ),
        (
            "Transparency by design",
            " — Approved projects can be published to the public portal. Leadership sees internal performance; citizens see what the county delivers.",
        ),
        (
            "Built for Kenyan county government",
            " — Departments, sub-counties, wards, villages, financial years, procurement stages, and reporting templates aligned to how counties actually work.",
        ),
    ]
    for bold, rest in capabilities:
        add_bullet(doc, bold, rest)

    add_heading(doc, "Why counties choose E-CIMES", level=2)
    reasons = [
        ("Proven in production", " — Deployed and operational for the County Government of Machakos, managing the full project portfolio across departments and wards."),
        ("Configurable per county", " — Branding, organisation structure, labels, and features adapt through county configuration—no rebuild required."),
        ("Secure, role-based access", " — Each user sees only what their role and organisation scope allow, from village administrators to the Governor's office."),
        ("Modern, maintainable stack", " — Web application (React), API (Node.js), PostgreSQL database, Docker deployment, and optional cloud or on-premise hosting."),
        ("Low infrastructure footprint", " — Runs on modest servers (from ~4 GB RAM for smaller deployments); scales as the county grows."),
    ]
    for bold, rest in reasons:
        add_bullet(doc, bold, rest)

    add_heading(doc, "Who uses it", level=2)
    users_table = doc.add_table(rows=8, cols=2)
    users_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    users_rows = [
        ("Role", "What they do in E-CIMES"),
        ("Governor / CEC / Chief Officers", "Summary statistics, finance dashboards, programme progress"),
        ("Village / Ward / Sub-County M&E", "Draft, review, and escalate monitoring visits"),
        ("Engineers & Co-Finance", "Payment certificates, progress photos, absorption tracking"),
        ("Procurement & project officers", "Budget-to-project conversion, contractor management"),
        ("Contractors", "Assigned-project portal for payments, photos, and files"),
        ("ICT & administrators", "User management, data import, workflow configuration"),
        ("Citizens", "Public portal—projects, maps, feedback (no login required)"),
    ]
    for r_idx, (col1, col2) in enumerate(users_rows):
        row = users_table.rows[r_idx]
        fill_table_cell(row.cells[0], col1, bold=(r_idx == 0))
        fill_table_cell(row.cells[1], col2, bold=(r_idx == 0))
        if r_idx == 0:
            style_table_header(row)
        elif r_idx % 2 == 0:
            for cell in row.cells:
                set_cell_shading(cell, TABLE_ALT_BG)

    doc.add_paragraph()

    add_heading(doc, "Implementation approach", level=2)
    add_body(
        doc,
        "E-CIMES is delivered as a turnkey county platform: environment setup, county configuration, data "
        "migration support, user provisioning, training materials (SOPs and guided help), and ongoing technical "
        "support. Typical rollout phases cover ICT setup, master data import, pilot departments, field mobile "
        "deployment, and public portal launch.",
    )

    add_heading(doc, "See it in action", level=2)
    p = doc.add_paragraph()
    add_run(p, "Live deployment: ", bold=True, size=10.5, color=TEXT_DARK)
    add_run(p, "https://cimes.machakos.go.ke", size=10.5, color=BRAND_ACCENT)
    set_paragraph_spacing(p, after=3)
    add_body(
        doc,
        "Public portal: County public investment dashboard (projects, regional breakdown, citizen feedback).",
        after=12,
    )

    closing = doc.add_paragraph()
    closing.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(closing, "E-CIMES", bold=True, size=12, color=BRAND_DARK)
    add_run(closing, " — ", size=12, color=TEXT_MUTED)
    add_run(closing, "Plan with clarity. Deliver with evidence. Account to citizens.", italic=True, size=12, color=BRAND_ACCENT)
    set_paragraph_spacing(closing, before=6, after=6)

    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(
        footer,
        "For demonstrations, technical briefings, or county deployment proposals, contact the E-CIMES implementation team.",
        italic=True,
        size=9.5,
        color=TEXT_MUTED,
    )

    return doc


def main():
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = build_document()
    doc.save(OUTPUT)
    print(f"Created: {OUTPUT}")


if __name__ == "__main__":
    main()
