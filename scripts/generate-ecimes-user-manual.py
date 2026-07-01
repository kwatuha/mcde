#!/usr/bin/env python3
"""
Generate printable E-CIMES User Manual (Word) from help-knowledge-base.json.
Same content source as Help & Support (/help-support) and the AI assistant manual.
"""

from __future__ import annotations

import json
import shutil
import subprocess
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
KB_PATH = ROOT / "frontend" / "src" / "data" / "help-knowledge-base.json"
LOGO_PATH = ROOT / "frontend" / "src" / "assets" / "gpris.png"
OUT_DOCX = ROOT / "docs" / "E-CIMES-User-Manual.docx"
OUT_PDF = ROOT / "docs" / "E-CIMES-User-Manual.pdf"

NAVY = RGBColor(0x0D, 0x47, 0xA1)
BLUE = RGBColor(0x15, 0x65, 0xC0)
GREEN = RGBColor(0x2E, 0x7D, 0x32)
SLATE = RGBColor(0x33, 0x33, 0x33)
HEADER_FILL = "0D47A1"
ALT_FILL = "E3F2FD"
MANUAL_VERSION = "1.0"
SYSTEM_FLOW = [
    "Planning setup",
    "Project registration",
    "Procurement",
    "Implementation tracking",
    "Monitoring & field data",
    "Finance & certificates",
    "Reporting & AI outputs",
    "Public transparency",
    "Administration & audit",
]
FIRST_CHECKS = [
    "Am I using the correct URL?",
    "Is my account active and approved?",
    "Do I have the right role and organisation scope?",
    "Are filters hiding the record I expect?",
    "For AI reports — am I on the correct dashboard first?",
    "For certificates — try QR scan or Finance → Verify Certificate.",
    "Did I save or confirm the action?",
]


def load_kb() -> dict:
    with open(KB_PATH, encoding="utf-8") as f:
        return json.load(f)


def shade_cell(cell, fill: str) -> None:
    sh = OxmlElement("w:shd")
    sh.set(qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(sh)


def set_table_borders(table, color: str = "0D47A1") -> None:
    tbl_pr = table._tbl.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        table._tbl.insert(0, tbl_pr)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:color"), color)
        borders.append(el)
    tbl_pr.append(borders)


def set_bottom_border(paragraph, color: str = "1565C0", size: int = 8) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def setup_doc(doc: Document) -> None:
    for sec in doc.sections:
        sec.top_margin = Inches(0.85)
        sec.bottom_margin = Inches(0.75)
        sec.left_margin = Inches(0.9)
        sec.right_margin = Inches(0.9)
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)


def add_run_para(doc, text, *, bold=False, size=11, align=None, color=None, italic=False, space_after=6):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.name = "Calibri"
    r.font.size = Pt(size)
    if color:
        r.font.color.rgb = color
    return p


def section_heading(doc, text, level=1):
    if level == 1:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(8)
        r = p.add_run(text)
        r.bold = True
        r.font.size = Pt(16)
        r.font.color.rgb = NAVY
        set_bottom_border(p)
        return p
    p = doc.add_heading(text, level=level)
    for r in p.runs:
        r.font.color.rgb = NAVY
        r.font.size = Pt(13 if level == 2 else 12)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(item, style="List Bullet")
        for r in p.runs:
            r.font.name = "Calibri"
            r.font.size = Pt(10.5)


def add_labeled_block(doc, label, text, label_color):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    lr = p.add_run(f"{label}: ")
    lr.bold = True
    lr.font.color.rgb = label_color
    lr.font.size = Pt(10.5)
    tr = p.add_run(text)
    tr.font.size = Pt(10.5)


def add_table(doc, headers, rows, col_widths=None):
    if not rows:
        return
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    set_table_borders(table)
    for i, h in enumerate(headers):
        c = table.rows[0].cells[i]
        c.text = h
        shade_cell(c, HEADER_FILL)
        for p in c.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.name = "Calibri"
                r.font.size = Pt(9)
                r.font.color.rgb = RGBColor(255, 255, 255)
    for ri, row in enumerate(rows):
        fill = ALT_FILL if ri % 2 == 1 else None
        for ci, val in enumerate(row):
            c = table.rows[ri + 1].cells[ci]
            c.text = str(val)
            if fill:
                shade_cell(c, fill)
            for p in c.paragraphs:
                for r in p.runs:
                    r.font.name = "Calibri"
                    r.font.size = Pt(9)
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)
    doc.add_paragraph()


def cover_page(doc: Document) -> None:
    if LOGO_PATH.is_file():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(LOGO_PATH), width=Inches(1.15))

    add_run_para(doc, "REPUBLIC OF KENYA", bold=True, size=12, align=WD_ALIGN_PARAGRAPH.CENTER, color=SLATE)
    add_run_para(
        doc,
        "COUNTY GOVERNMENT OF MACHAKOS",
        bold=True,
        size=14,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        color=NAVY,
    )
    add_run_para(
        doc,
        "Department of Finance and Economic Planning",
        size=11,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        color=SLATE,
        space_after=18,
    )

    add_run_para(
        doc,
        "E-CIMES User Manual",
        bold=True,
        size=24,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        color=NAVY,
    )
    add_run_para(
        doc,
        "Electronic County Integrated Monitoring and Evaluation System",
        size=13,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        color=SLATE,
    )
    add_run_para(
        doc,
        "Printable staff guide — workflows, dashboards, finance, monitoring,\n"
        "AI assistant, mobile field collection, reports & troubleshooting",
        size=11,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        italic=True,
        space_after=24,
    )

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for line in [
        f"Manual version: {MANUAL_VERSION}",
        f"Published: {date.today().strftime('%d %B %Y')}",
        "Source: Help & Support (in-system) — /help-support",
        "For internal county staff use",
    ]:
        r = meta.add_run(line + "\n")
        r.font.size = Pt(10)
        r.font.color.rgb = SLATE

    doc.add_page_break()


def toc_page(doc: Document) -> None:
    section_heading(doc, "Contents")
    entries = [
        "1. About this manual",
        "2. Logical flow of the system",
        "3. First things to check",
        "4. Role-based guidance",
        "5. Quick task guide",
        "6. Module-by-module user guide",
        "7. Dashboard reference",
        "8. Key workflows (certificates, AI, mobile, reports)",
        "9. AI Assistant & professional reports",
        "10. Good data practice",
        "11. Troubleshooting guide",
        "12. Before contacting ICT support",
    ]
    for e in entries:
        add_run_para(doc, e, size=10.5, space_after=3)
    doc.add_page_break()


def build_manual(kb: dict) -> Document:
    doc = Document()
    setup_doc(doc)
    cover_page(doc)
    toc_page(doc)

    # 1 About
    section_heading(doc, "1. About this manual")
    add_run_para(
        doc,
        "This manual is the printable edition of the in-system Help & Support guide for E-CIMES "
        "(Electronic County Integrated Monitoring and Evaluation System). It supports county staff "
        "in planning, project management, procurement, monitoring, finance, reporting, and public "
        "transparency workflows.",
        space_after=8,
    )
    add_labeled_block(
        doc,
        "Important",
        "The system is role-based. If a screen, button, project, or report is missing, confirm your "
        "role, permissions, organisation scope, and active filters before escalating to ICT.",
        NAVY,
    )
    add_labeled_block(
        doc,
        "AI Assistant",
        "The sparkle button (bottom-right) uses this manual for navigation questions and live data "
        "from the page you are viewing for reports and summaries.",
        BLUE,
    )

    # 2 Logical flow
    section_heading(doc, "2. Logical flow of the system")
    add_run_para(
        doc,
        "Most county work in E-CIMES follows this path from planning to reporting and public accountability:",
        space_after=8,
    )
    for i, step in enumerate(SYSTEM_FLOW, 1):
        p = doc.add_paragraph(style="List Number")
        r = p.add_run(f"Step {i} — {step}")
        r.bold = True
        r.font.size = Pt(10.5)

    # 3 First checks
    section_heading(doc, "3. First things to check")
    add_run_para(doc, "These checks solve many support issues before escalation:", space_after=6)
    add_bullets(doc, FIRST_CHECKS)

    # 4 Role guidance
    section_heading(doc, "4. Role-based guidance")
    add_table(
        doc,
        ["Role / user group", "Primary responsibility"],
        kb.get("roleGuidance", []),
        col_widths=[1.6, 4.6],
    )

    # 5 Quick tasks
    section_heading(doc, "5. Quick task guide")
    add_run_para(doc, "Start here when you know what you want to do but are unsure where to go:", space_after=6)
    add_table(
        doc,
        ["Task", "Where to go", "What to do"],
        kb.get("quickTasks", []),
        col_widths=[1.4, 1.8, 3.0],
    )

    doc.add_page_break()

    # 6 Module guides
    section_heading(doc, "6. Module-by-module user guide")
    add_run_para(
        doc,
        "Each module below lists purpose, menu route, step-by-step use, and good practice.",
        space_after=10,
    )
    for guide in kb.get("moduleGuides", []):
        section_heading(doc, guide.get("title", "Module"), level=2)
        add_labeled_block(doc, "Audience", guide.get("audience", "All users"), BLUE)
        add_labeled_block(doc, "Route", guide.get("route", "—"), GREEN)
        add_labeled_block(doc, "Purpose", guide.get("purpose", ""), SLATE)
        doc.add_paragraph()
        add_run_para(doc, "How to use this area", bold=True, size=11, color=NAVY, space_after=4)
        add_bullets(doc, guide.get("steps", []))
        add_run_para(doc, "Good practice", bold=True, size=11, color=NAVY, space_after=4)
        add_bullets(doc, guide.get("tips", []))
        doc.add_paragraph()

    doc.add_page_break()

    # 7 Dashboard reference
    section_heading(doc, "7. Dashboard reference")
    dash_rows = []
    for d in kb.get("dashboardGuides", []):
        dash_rows.append([
            d.get("title", ""),
            d.get("menuPath", ""),
            d.get("purpose", ""),
            d.get("aiReportHint", "—"),
        ])
    add_table(
        doc,
        ["Dashboard / report", "Menu path", "Purpose", "AI report hint"],
        dash_rows,
        col_widths=[1.3, 1.5, 2.2, 1.2],
    )

    # 8 Navigation topics
    section_heading(doc, "8. Key workflows")
    for topic in kb.get("navigationTopics", []):
        section_heading(doc, topic.get("title", ""), level=2)
        add_labeled_block(doc, "Menu", topic.get("menuPath", ""), GREEN)
        add_labeled_block(doc, "Summary", topic.get("summary", ""), SLATE)
        if topic.get("steps"):
            add_run_para(doc, "Steps", bold=True, size=11, color=NAVY, space_after=4)
            add_bullets(doc, topic["steps"])
        if topic.get("tips"):
            add_run_para(doc, "Tips", bold=True, size=11, color=NAVY, space_after=4)
            add_bullets(doc, topic["tips"])
        doc.add_paragraph()

    doc.add_page_break()

    # 9 AI guide
    ai = kb.get("aiAssistantGuide", {})
    section_heading(doc, "9. AI Assistant & professional reports")
    add_labeled_block(doc, "Summary", ai.get("summary", ""), SLATE)
    add_run_para(doc, "Steps", bold=True, size=11, color=NAVY, space_after=4)
    add_bullets(doc, ai.get("steps", []))
    add_run_para(doc, "Tips", bold=True, size=11, color=NAVY, space_after=4)
    add_bullets(doc, ai.get("tips", []))

    # 10 Good practice
    section_heading(doc, "10. Good data practice")
    add_bullets(doc, kb.get("goodPracticeItems", []))

    # 11 Troubleshooting
    section_heading(doc, "11. Troubleshooting guide")
    add_table(
        doc,
        ["Issue", "Likely cause", "First action"],
        kb.get("troubleshootingRows", []),
        col_widths=[1.3, 1.8, 3.1],
    )
    add_labeled_block(
        doc,
        "Warning",
        "Do not send screenshots containing passwords, OTP codes, private tokens, or confidential "
        "personal data. Mask sensitive information before sharing support evidence.",
        NAVY,
    )

    # 12 Support checklist
    section_heading(doc, "12. Before contacting ICT support")
    add_run_para(doc, "Include these details so support can reproduce and resolve the issue quickly:", space_after=6)
    add_bullets(doc, kb.get("supportChecklist", []))

    # Footer
    doc.add_paragraph()
    add_run_para(
        doc,
        "— End of manual —",
        bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        color=NAVY,
        space_after=6,
    )
    add_run_para(
        doc,
        "Open the live guide anytime: sign in → three-dot menu (top right) → Help & Support.\n"
        "County Government of Machakos · E-CIMES",
        size=9,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        italic=True,
        color=SLATE,
    )

    return doc


def export_pdf(docx_path: Path, pdf_path: Path) -> bool:
    soffice = shutil.which("libreoffice") or shutil.which("soffice")
    if not soffice:
        return False
    try:
        subprocess.run(
            [soffice, "--headless", "--convert-to", "pdf", str(docx_path), "--outdir", str(pdf_path.parent)],
            check=True,
            capture_output=True,
            timeout=120,
        )
    except (subprocess.CalledProcessError, subprocess.TimeoutExpired):
        return False
    return pdf_path.is_file()


def main():
    if not KB_PATH.is_file():
        raise SystemExit(f"Knowledge base not found: {KB_PATH}")
    OUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    kb = load_kb()
    doc = build_manual(kb)
    doc.save(OUT_DOCX)
    print(f"Wrote {OUT_DOCX}")
    print(f"  Modules: {len(kb.get('moduleGuides', []))}")
    print(f"  Quick tasks: {len(kb.get('quickTasks', []))}")
    if export_pdf(OUT_DOCX, OUT_PDF):
        print(f"Wrote {OUT_PDF}")
    else:
        print("  PDF: install LibreOffice or open the .docx → File → Export as PDF.")
    print("  Print: open either file → File → Print.")


if __name__ == "__main__":
    main()
